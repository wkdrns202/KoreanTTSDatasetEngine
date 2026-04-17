"""
Korean TTS Research Progress Report Generator
Generates a comprehensive .docx report covering Data Collection, Cleansing, and Model Training.
"""

import json
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Paths ──────────────────────────────────────────────────────────────
BASE = Path(__file__).resolve().parent.parent
OUTPUT = BASE / "docs" / "TTS_Research_Progress_Report.docx"

# ── Constants ──────────────────────────────────────────────────────────
FONT_KOREAN = "Malgun Gothic"
FONT_CODE = "Consolas"
FONT_SIZE_BODY = Pt(10.5)
FONT_SIZE_SMALL = Pt(9)
FONT_SIZE_TABLE = Pt(9.5)

COLOR_PRIMARY = RGBColor(0x1A, 0x56, 0xDB)    # Blue
COLOR_HEADING = RGBColor(0x1F, 0x29, 0x37)     # Dark navy
COLOR_ACCENT = RGBColor(0x2E, 0x7D, 0x32)      # Green
COLOR_GRAY = RGBColor(0x66, 0x66, 0x66)
COLOR_TABLE_HEADER = "1A56DB"
COLOR_TABLE_ALT = "F0F4FA"
COLOR_PASS = RGBColor(0x2E, 0x7D, 0x32)
COLOR_FAIL = RGBColor(0xC6, 0x28, 0x28)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, font_name=FONT_KOREAN, size=FONT_SIZE_BODY, bold=False, color=None, italic=False):
    """Configure run font properties."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_styled_paragraph(doc, text, font_name=FONT_KOREAN, size=FONT_SIZE_BODY,
                         bold=False, color=None, alignment=None, space_after=Pt(6),
                         space_before=Pt(0), italic=False):
    """Add a paragraph with consistent styling."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold, color, italic)
    return p


def add_bullet(doc, text, level=0, bold_prefix="", font_size=FONT_SIZE_BODY):
    """Add a bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_run_font(run, size=font_size, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=font_size)
    return p


def add_code_block(doc, text):
    """Add a code-style paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_run_font(run, FONT_CODE, Pt(8.5))
    # Light gray background via shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    pPr.append(shading)
    return p


def create_table(doc, headers, rows, col_widths=None):
    """Create a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=FONT_SIZE_TABLE, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, COLOR_TABLE_HEADER)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_run_font(run, size=FONT_SIZE_TABLE)
            # Alternate row shading
            if r_idx % 2 == 1:
                set_cell_shading(cell, COLOR_TABLE_ALT)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = width

    doc.add_paragraph()  # spacing after table
    return table


def add_heading(doc, text, level=1):
    """Add heading with Korean font."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_KOREAN
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KOREAN)
    return h


# ══════════════════════════════════════════════════════════════════════
#  DOCUMENT SECTIONS
# ══════════════════════════════════════════════════════════════════════

def create_cover_page(doc):
    """Create professional cover page."""
    # Add spacing
    for _ in range(6):
        doc.add_paragraph()

    # Title (Korean)
    add_styled_paragraph(
        doc, "한국어 TTS 연구 진행 보고서",
        size=Pt(28), bold=True, color=COLOR_HEADING,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12)
    )
    # Subtitle (English)
    add_styled_paragraph(
        doc, "Korean Text-to-Speech Research Progress Report",
        size=Pt(16), color=COLOR_PRIMARY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(48)
    )

    # Divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    set_run_font(run, size=Pt(12), color=COLOR_PRIMARY)

    doc.add_paragraph()

    # Metadata
    meta_items = [
        ("프로젝트", "한국어 TTS 음성합성 파이프라인 구축"),
        ("작성일", datetime.now().strftime("%Y년 %m월 %d일")),
        ("작성자", "AI Research Team"),
        ("문서 버전", "v1.0"),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{label}: ")
        set_run_font(run, size=Pt(12), bold=True, color=COLOR_HEADING)
        run = p.add_run(value)
        set_run_font(run, size=Pt(12), color=COLOR_GRAY)

    doc.add_page_break()


def section_1_overview(doc):
    """1. Project Overview."""
    add_heading(doc, "1. 프로젝트 개요 (Project Overview)", level=1)

    add_heading(doc, "1.1 목적 및 배경 (Purpose & Background)", level=2)
    add_styled_paragraph(
        doc,
        "본 프로젝트는 고품질 한국어 TTS(Text-to-Speech) 음성합성 시스템 구축을 목표로 합니다. "
        "전문 성우의 녹음 데이터를 기반으로, 자동화된 데이터 정제 파이프라인을 통해 학습용 데이터셋을 "
        "생성하고, 최신 TTS 모델(F5-TTS)을 fine-tuning하여 자연스러운 한국어 음성을 합성합니다."
    )
    add_styled_paragraph(
        doc,
        "기존 한국어 TTS 데이터셋 구축은 수작업 레이블링에 의존하여 비용과 시간이 많이 소요되었습니다. "
        "본 연구에서는 OpenAI Whisper ASR을 활용한 자동 정렬(alignment) 기법을 통해 "
        "이 과정을 자동화하고, 반복적 R&D 사이클을 통해 95% 이상의 정확도를 달성하였습니다."
    )

    add_heading(doc, "1.2 전체 파이프라인 구조 (End-to-End Pipeline Architecture)", level=2)
    add_styled_paragraph(
        doc,
        "프로젝트는 크게 세 단계로 구성됩니다:",
        bold=True
    )

    pipeline_flow = (
        "녹음 (Record)  →  데이터 정제 (Cleanse)  →  모델 학습 (Train)  →  추론 (Inference)"
    )
    add_styled_paragraph(
        doc, pipeline_flow,
        size=Pt(12), bold=True, color=COLOR_PRIMARY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(12), space_after=Pt(12)
    )

    steps = [
        ("1단계 - 데이터 수집: ", "전문 성우 녹음, 19개 오디오 파일(~7.1GB), 5개 스크립트(~4,845 문장)"),
        ("2단계 - 데이터 정제: ", "Whisper ASR 기반 자동 정렬, 6-Stage 파이프라인, 7회 반복 R&D"),
        ("3단계 - 모델 학습: ", "F5-TTS v1 Base fine-tuning, RunPod RTX 4090, Phase 1 완료(~22K steps)"),
        ("4단계 - 추론/배포: ", "CLI, Gradio UI, Python API 지원, 체크포인트 경량화(5.1GB → 1.3GB)"),
    ]
    for prefix, text in steps:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading(doc, "1.3 기술 스택 요약 (Technology Stack Summary)", level=2)

    tech_rows = [
        ["분류", "기술", "세부 사항"],
    ]
    tech_data = [
        ["음성 인식 (ASR)", "OpenAI Whisper", "medium (정렬), large (검증), language='ko'"],
        ["TTS 모델", "F5-TTS v1 Base", "Flow-matching 기반, vocab 확장 (2,545→3,391)"],
        ["학습 인프라", "RunPod Cloud GPU", "RTX 4090 (24GB VRAM), $0.59/hr"],
        ["개발 언어", "Python 3.11", "pytorch, numpy, soundfile, python-Levenshtein"],
        ["오디오 처리", "soundfile / numpy", "48kHz, 24-bit mono WAV"],
        ["버전 관리", "Git / GitHub", "TTSDataSetCleanser, korean-tts-research"],
        ["모델 저장소", "HuggingFace Hub", "DavidVita/Korean_Chaboon, F5TTS-Korean-Training"],
        ["GPU (로컬)", "RTX 3060 Ti", "8GB VRAM, 데이터 정제용"],
        ["GPU (클라우드)", "RTX 4090", "24GB VRAM, 모델 학습용"],
    ]
    create_table(doc,
                 ["분류", "기술", "세부 사항"],
                 tech_data)


def section_2_data_collection(doc):
    """2. Data Collection."""
    add_heading(doc, "2. 데이터 수집 (Data Collection / Recording)", level=1)

    add_heading(doc, "2.1 원본 데이터 사양 (Raw Data Specifications)", level=2)

    create_table(doc,
                 ["항목", "사양"],
                 [
                     ["오디오 파일 수", "19개"],
                     ["총 용량", "~7.1 GB"],
                     ["샘플링 레이트", "48,000 Hz"],
                     ["비트 깊이", "24-bit"],
                     ["채널", "Mono"],
                     ["포맷", "WAV (PCM)"],
                     ["녹음자", "전문 성우 1인 (여성)"],
                 ])

    add_heading(doc, "2.2 스크립트 구성 (Script Composition)", level=2)
    add_styled_paragraph(
        doc,
        "5개의 스크립트로 구성되며, 각 스크립트는 번호|텍스트 형식의 한국어 문장을 포함합니다."
    )

    create_table(doc,
                 ["스크립트", "문장 수", "오디오 파일 수", "내용 특성"],
                 [
                     ["Script 1", "300", "3", "문학 작품 낭독 (이상한 나라의 앨리스 등)"],
                     ["Script 2", "1,644", "6", "IT/프로그래밍 교육 콘텐츠"],
                     ["Script 3", "878", "2", "문학 텍스트 (고전/현대문학)"],
                     ["Script 4", "1,005", "5", "일반 교양 콘텐츠"],
                     ["Script 5", "1,018*", "3", "내러티브/서사 콘텐츠"],
                 ])

    add_styled_paragraph(
        doc,
        "* Script 5는 541번 라인까지만 오디오가 존재하며, 542~1,018번은 미녹음 상태입니다.",
        size=FONT_SIZE_SMALL, italic=True, color=COLOR_GRAY
    )

    add_styled_paragraph(
        doc,
        "총 타겟 문장: ~4,845줄 (오디오가 존재하는 구간 기준: ~4,228줄)"
    )

    add_heading(doc, "2.3 녹음 환경 및 포맷 (Recording Environment & Format)", level=2)
    bullets = [
        ("녹음 포맷: ", "48kHz / 24-bit / Mono WAV (무손실 PCM)"),
        ("파일 명명 규칙: ", "Script_{N}_{Start}-{End}.wav (스크립트 번호_시작라인-끝라인)"),
        ("인코딩: ", "UTF-8-sig (BOM 포함), 일부 파일 CP949/EUC-KR"),
        ("텍스트 형식: ", "idx|text (파이프 구분자, 순번|한국어 문장)"),
    ]
    for prefix, text in bullets:
        add_bullet(doc, text, bold_prefix=prefix)


def section_3_cleansing_pipeline(doc):
    """3. Data Cleansing Pipeline."""
    add_heading(doc, "3. 데이터 정제 파이프라인 (Data Cleansing Pipeline)", level=1)

    add_styled_paragraph(
        doc,
        "데이터 정제는 본 프로젝트의 핵심 R&D 영역입니다. 6-Stage 반복 파이프라인을 설계하여 "
        "자동화된 정렬, 검증, 진단, 개선 사이클을 구현하였으며, 7회의 반복(iteration)을 통해 "
        "최초 34.7%에서 최종 95.48%의 정확도를 달성하였습니다."
    )

    # 3.1 Pipeline Architecture
    add_heading(doc, "3.1 파이프라인 아키텍처 (Pipeline Architecture)", level=2)
    add_styled_paragraph(
        doc,
        "파이프라인은 6개의 Stage로 구성된 반복형(iterative) 구조입니다:"
    )

    stages_data = [
        ["Stage 1", "Align & Split", "Whisper ASR로 음성 전사 후 스크립트와 정렬, 개별 WAV로 분할"],
        ["Stage 2", "Clean & Post-process", "Zero-crossing snap, fade, 오디오 envelope 적용"],
        ["Stage 3", "Validate", "분할된 WAV을 재전사(re-transcribe)하여 원문과 비교 검증"],
        ["Stage 4", "Evaluate", "전체 메트릭 집계, R1~R6 요구사항 평가"],
        ["Stage 5", "Finalize", "95% 이상 달성 시 최종 데이터셋 확정"],
        ["Stage 6", "Diagnose & Improve", "95% 미만 시 실패 분석, 파라미터 조정 후 Stage 1로 복귀"],
    ]
    create_table(doc,
                 ["Stage", "이름", "설명"],
                 stages_data)

    add_styled_paragraph(
        doc,
        "Stage 6에서 Stage 1으로 되돌아가는 반복 구조를 통해, 각 iteration마다 실패 원인을 "
        "분석하고 파라미터를 조정하여 점진적으로 품질을 개선합니다. "
        "최대 10회 반복 후 수동 개입을 요청하는 안전장치가 포함되어 있습니다."
    )

    # 3.2 Stage 1: Align & Split
    add_heading(doc, "3.2 Stage 1: Align & Split — Whisper ASR 정렬", level=2)
    add_styled_paragraph(
        doc,
        "핵심 정렬 알고리즘은 Whisper medium 모델을 사용하여 장시간 오디오를 전사한 후, "
        "각 전사 세그먼트를 원본 스크립트 라인과 매칭합니다."
    )

    add_styled_paragraph(doc, "주요 알고리즘 구성요소:", bold=True, space_before=Pt(8))

    algo_items = [
        ("Forward-only Search (전방 탐색): ",
         "현재 스크립트 위치에서 25줄 전방만 탐색합니다. 후방 탐색은 한국어에서 "
         "중복 매칭(duplicate match)을 유발하므로 사용하지 않습니다."),
        ("Segment Merging (세그먼트 병합): ",
         "Whisper가 하나의 문장을 여러 세그먼트로 분할하는 경우가 많으므로, "
         "1~5개의 연속 세그먼트를 결합하여 최적 매칭을 시도합니다."),
        ("Skip Penalty (건너뛰기 패널티): ",
         "현재 위치에서 먼 라인에 매칭될수록 0.01/줄의 패널티를 부과하여 "
         "정렬 순서를 유지합니다."),
        ("Match Confirmation (매칭 확인): ",
         "유사도 0.70 미만의 경계선 매칭에 대해 추가 검증을 수행하여 "
         "false positive를 방지합니다."),
        ("Re-sync (재동기화): ",
         "10회 연속 매칭 실패 시, 75줄 범위에서 threshold 0.35로 "
         "재동기화를 시도합니다."),
        ("Gap-aware Padding: ",
         "세그먼트 간 실제 silence gap이 있을 때만 padding을 적용하고, "
         "back-to-back 경계에서는 zero-pad로 sentence bleeding을 방지합니다."),
    ]
    for prefix, text in algo_items:
        add_bullet(doc, text, bold_prefix=prefix)

    # 3.3 Stage 2: Clean & Post-process
    add_heading(doc, "3.3 Stage 2: Clean & Post-process — 오디오 후처리", level=2)
    add_styled_paragraph(
        doc,
        "각 분할된 WAV 파일에 대해 다음 후처리를 순차 적용합니다:"
    )

    postproc = [
        ("Zero-crossing Snap: ", "절단점을 ±10ms 내 가장 가까운 영점 교차점으로 이동하여 클릭 방지"),
        ("Fade In/Out: ", "Raised-cosine envelope로 10ms fade-in/fade-out 적용"),
        ("Silence Trimming: ", "RMS < -40dB 기준으로 선행/후행 silence 제거"),
        ("R6 Audio Envelope: ", "TTS 학습을 위한 표준 오디오 envelope 적용"),
        ("Peak Normalization: ", "음성 구간만 -1dB로 정규화 (padding silence 제외)"),
    ]
    for prefix, text in postproc:
        add_bullet(doc, text, bold_prefix=prefix)

    add_styled_paragraph(doc, "R6 Audio Envelope 구성:", bold=True, space_before=Pt(8))
    add_styled_paragraph(
        doc,
        "모든 WAV 파일은 [400ms Pre-attack Silence] + [음성] + [730ms Tail Silence] 구조를 갖습니다. "
        "Onset Safety Margin 30ms로 자음 공격(consonant attack)을 보호하고, "
        "Offset Safety Margin 80ms로 자연스러운 음성 감쇠를 유지합니다."
    )

    create_table(doc,
                 ["파라미터", "값", "설명"],
                 [
                     ["Pre-attack Silence", "400ms", "음성 시작 전 보장된 무음 구간"],
                     ["Tail Silence", "730ms", "음성 종료 후 보장된 무음 구간"],
                     ["Onset Safety Margin", "30ms", "감지된 음성 시작점 이전 안전 여백"],
                     ["Offset Safety Margin", "80ms", "감지된 음성 종료점 이후 안전 여백"],
                     ["AUDIO_PAD_MS", "100ms", "오디오 추출 시 기본 padding"],
                     ["Fade Duration", "10ms", "Fade-in/fade-out 길이"],
                 ])

    # 3.4 Stage 3-4: Validate & Evaluate
    add_heading(doc, "3.4 Stage 3-4: Validate & Evaluate — 검증 및 평가", level=2)
    add_styled_paragraph(
        doc,
        "분할된 각 WAV 파일을 Whisper로 재전사(re-transcribe)한 후, "
        "원본 스크립트와의 유사도를 계산하여 정렬 품질을 평가합니다."
    )

    add_styled_paragraph(doc, "Tiered Evaluation 전략:", bold=True, space_before=Pt(8))
    add_styled_paragraph(
        doc,
        "단일 모델 평가의 한계를 극복하기 위해 2-tier 평가 전략을 도입하였습니다."
    )

    eval_data = [
        ["Tier 1 (Primary)", "Whisper medium", "GT-prompted + unprompted, best-of-2", "~1초/항목"],
        ["Tier 2 (Recovery)", "Whisper large", "Tier 1 실패 항목 중 sim ≥ 0.50만 재평가", "~20초/항목"],
    ]
    create_table(doc,
                 ["단계", "모델", "전략", "처리 속도"],
                 eval_data)

    add_styled_paragraph(
        doc,
        "Tier 2 대형 모델은 Tier 1 실패 항목의 약 15%를 추가 복구합니다. "
        "RTX 3060 Ti (8GB VRAM) 환경에서 두 모델을 순차적으로 로드해야 하며, "
        "모델 전환 시 반드시 CPU로 이동 후 gc.collect()를 수행해야 CUDA OOM을 방지할 수 있습니다."
    )

    add_styled_paragraph(doc, "평가 전처리:", bold=True, space_before=Pt(8))
    bullets_eval = [
        ("Envelope Strip: ", "Whisper 평가 전 R6 envelope을 제거하여 첫 음절 drop 방지 "
         "(EVAL_STRIP_LEAD_MS=350, EVAL_STRIP_TAIL_MS=700)"),
        ("텍스트 정규화: ", "구두점 제거, 한글+영숫자만 유지, Unicode NFC 정규화"),
        ("유사도 계산: ", "CER = levenshtein_distance / max(len(gt), len(whisper)), similarity = 1 - CER"),
        ("Timeout: ", "파일당 120초 제한 (무한 디코딩 루프 방지)"),
    ]
    for prefix, text in bullets_eval:
        add_bullet(doc, text, bold_prefix=prefix)

    # 3.5 Stage 5-6
    add_heading(doc, "3.5 Stage 5-6: Finalize / Diagnose — 확정 및 진단", level=2)
    add_styled_paragraph(
        doc,
        "Stage 4에서 모든 요구사항(R1~R6)이 95% 이상을 달성하면 Stage 5에서 데이터셋을 확정합니다. "
        "미달 시 Stage 6에서 실패 원인을 분석하고 파라미터를 조정하여 Stage 1부터 재실행합니다."
    )

    add_styled_paragraph(doc, "실패 유형 분류 체계:", bold=True, space_before=Pt(8))

    fail_types = [
        ["Type A", "Alignment Shift", "전사 텍스트가 정확하지만 잘못된 스크립트 라인에 매칭"],
        ["Type B", "Merge/Split Error", "두 문장이 하나의 WAV로 병합되거나 하나가 분할됨"],
        ["Type C", "Boundary Noise", "인접 문장의 음성이 경계에서 bleed"],
        ["Type D", "Whisper Error", "Whisper가 올바른 오디오를 잘못 인식 (인식 한계)"],
        ["Type E", "Script Mismatch", "스크립트 텍스트 자체의 오류"],
        ["Type F", "Envelope Violation", "Pre-attack 또는 Tail silence 미달"],
    ]
    create_table(doc,
                 ["유형", "이름", "설명"],
                 fail_types)

    # 3.6 Key Parameters
    add_heading(doc, "3.6 핵심 파라미터 (Key Parameters)", level=2)
    add_styled_paragraph(
        doc,
        "7회의 반복 R&D를 통해 최적화된 최종 파라미터 값입니다."
    )

    params_data = [
        ["MODEL_SIZE", '"medium"', "Whisper 모델 크기 (Stage 1 정렬용)"],
        ["SEG_SEARCH_WINDOW", "25", "전방 탐색 범위 (줄 수)"],
        ["SKIP_PENALTY", "0.01", "건너뛰기 패널티 (줄당)"],
        ["MATCH_THRESHOLD", "0.50", "최소 조정 유사도 임계값"],
        ["CONSEC_FAIL_LIMIT", "10", "재동기화 트리거 연속 실패 횟수"],
        ["MAX_MERGE", "5", "최대 연속 세그먼트 병합 수"],
        ["AUDIO_PAD_MS", "100", "기본 오디오 padding (ms)"],
        ["MIN_GAP_FOR_PAD_MS", "20", "실제 padding 적용 최소 gap (ms)"],
        ["FADE_MS", "10", "Fade-in/out 길이 (ms)"],
        ["PREATTACK_SILENCE_MS", "400", "Pre-attack silence 길이 (ms)"],
        ["TAIL_SILENCE_MS", "730", "Tail silence 길이 (ms)"],
        ["ONSET_SAFETY_MS", "30", "음성 onset 안전 여백 (ms)"],
        ["OFFSET_SAFETY_MS", "80", "음성 offset 안전 여백 (ms)"],
        ["Re-sync Window", "75줄, threshold 0.35", "재동기화 탐색 범위 및 임계값"],
    ]
    create_table(doc,
                 ["파라미터", "값", "설명"],
                 params_data)

    # 3.7 Iteration History
    add_heading(doc, "3.7 R&D 반복 이력 (Iteration History)", level=2)
    add_styled_paragraph(
        doc,
        "총 7회의 R&D 반복을 수행하여 R1 정확도를 34.7%에서 95.48%로 개선하였습니다. "
        "아래 표는 각 iteration의 핵심 변경사항과 결과를 요약합니다."
    )

    iter_data = [
        ["1", "2026-02-08", "34.7%", "초기 실행, MATCH_THRESHOLD=0.25",
         "Type A 1,757건 (alignment shift 지배적)"],
        ["2", "2026-02-08", "64.3%", "MATCH_THRESHOLD 0.50, GT-prompting 도입",
         "GT-prompting으로 R1 거의 2배 향상"],
        ["3", "2026-02-09", "69.7%", "Tiered evaluation (medium+large) 도입",
         "Tier 2 large 모델로 +5.4% 복구"],
        ["4", "2026-02-09", "~72.0%", "Match confirmation 도입 (유일한 생존 개선)",
         "다변수 변경의 위험성 학습, 1개씩 변경 원칙 확립"],
        ["5", "2026-02-09", "67.8%", "Inline verification 시도 (실패, 역효과)",
         "검증이 정렬 흐름을 방해 → REVERT"],
        ["6", "2026-02-09", "—", "Whisper large 평가 전환, inline verification 제거",
         "평가 모델 업그레이드 (medium→large)"],
        ["7", "2026-02-20", "95.48%", "Audio bleed 수정 (PAD 100→50, OFFSET 100→80)",
         "파라미터 regression 발견 및 수정, metadata 버그 수정"],
    ]
    create_table(doc,
                 ["Iter", "날짜", "R1", "핵심 변경사항", "결과 및 교훈"],
                 iter_data)

    # R1 progress summary
    add_styled_paragraph(
        doc,
        "R1 정확도 추이: 34.7% → 64.3% → 69.7% → 72.0% → 67.8%(↓) → — → 95.48%",
        bold=True, color=COLOR_PRIMARY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(8), space_after=Pt(12)
    )

    # 3.8 Final Quality Metrics
    add_heading(doc, "3.8 최종 품질 지표 (Final Quality Metrics)", level=2)
    add_styled_paragraph(
        doc,
        "Iteration 7 완료 후 최종 평가 결과입니다. 6개 요구사항 모두 기준을 충족하였습니다.",
        space_after=Pt(8)
    )

    # R1-R6 table with PASS/FAIL coloring
    req_table = doc.add_table(rows=7, cols=4)
    req_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    req_table.style = "Table Grid"
    req_headers = ["요구사항", "기준", "달성 점수", "판정"]
    for i, h in enumerate(req_headers):
        cell = req_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=FONT_SIZE_TABLE, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, COLOR_TABLE_HEADER)

    req_data = [
        ["R1 (Alignment Accuracy)", ">= 95%", "95.48%", "PASS"],
        ["R2 (Boundary Noise Clean)", ">= 95%", "100.0%", "PASS"],
        ["R3 (Combined Pass Rate)", ">= 95%", "95.4%", "PASS"],
        ["R4 (Metadata Integrity)", "Complete", "4,196/4,200 (99.9%)", "PASS"],
        ["R5 (Reproducibility)", "TRUE", "TRUE", "PASS"],
        ["R6 (Audio Envelope)", ">= 95%", "99.93%", "PASS"],
    ]
    for r_idx, row_data in enumerate(req_data):
        for c_idx, cell_text in enumerate(row_data):
            cell = req_table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            if c_idx == 3:  # PASS/FAIL column
                color = COLOR_PASS if cell_text == "PASS" else COLOR_FAIL
                set_run_font(run, size=FONT_SIZE_TABLE, bold=True, color=color)
            else:
                set_run_font(run, size=FONT_SIZE_TABLE)
            if r_idx % 2 == 1:
                set_cell_shading(cell, COLOR_TABLE_ALT)

    doc.add_paragraph()  # spacing

    add_styled_paragraph(doc, "R6 Envelope 상세 통계:", bold=True, space_before=Pt(4))
    create_table(doc,
                 ["항목", "최소값", "최대값", "평균"],
                 [
                     ["Pre-attack Silence", "400.0ms", "440.0ms", "409.31ms"],
                     ["Tail Silence", "724.35ms", "1,025.23ms", "737.99ms"],
                 ])


def section_4_final_dataset(doc):
    """4. Final Dataset."""
    add_heading(doc, "4. 최종 데이터셋 (Final Dataset)", level=1)

    add_heading(doc, "4.1 데이터셋 통계 (Dataset Statistics)", level=2)

    create_table(doc,
                 ["항목", "수치"],
                 [
                     ["script.txt 엔트리 수", "4,196"],
                     ["WAV 파일 수", "4,200"],
                     ["Orphan WAV (메타데이터 없음)", "4"],
                     ["검역(Quarantine) 파일", "74 (sim < 0.80)"],
                     ["오디오 포맷", "48kHz / 24-bit / Mono WAV"],
                     ["총 오디오 시간", "~8.18시간"],
                 ])

    add_heading(doc, "4.2 스크립트별 품질 분포 (Per-Script Quality)", level=2)

    create_table(doc,
                 ["스크립트", "총 세그먼트", "통과", "통과율"],
                 [
                     ["Script 1", "298", "294", "98.66%"],
                     ["Script 2", "1,254", "1,146", "91.39%"],
                     ["Script 3", "865", "841", "97.23%"],
                     ["Script 4", "990", "949", "95.86%"],
                     ["Script 5", "792", "776", "97.98%"],
                     ["전체 합계", "4,199", "4,006", "95.40%"],
                 ])

    add_styled_paragraph(
        doc,
        "Script 2가 91.39%로 가장 낮은 통과율을 보이는데, 이는 IT/프로그래밍 용어 "
        "(영어 혼용, 약어, 기술 용어)가 많아 Whisper 인식률이 떨어지기 때문입니다.",
        space_before=Pt(4)
    )

    add_heading(doc, "4.3 실패 분석 (Failure Analysis)", level=2)
    add_styled_paragraph(
        doc,
        "총 193건의 실패 세그먼트 중, Type D (Whisper 인식 오류)가 95.3%로 절대 다수를 차지합니다. "
        "이는 정렬 알고리즘의 문제가 아닌 Whisper 모델의 한국어 인식 한계에 기인합니다."
    )

    create_table(doc,
                 ["실패 유형", "건수", "비율", "설명"],
                 [
                     ["Type A (Alignment Shift)", "6", "3.1%", "잘못된 스크립트 라인 매칭"],
                     ["Type B (Merge/Split)", "0", "0%", "세그먼트 병합/분할 오류"],
                     ["Type C (Boundary Noise)", "0", "0%", "경계 음성 bleeding"],
                     ["Type D (Whisper Error)", "184", "95.3%", "Whisper 인식 한계"],
                     ["Type E (Script Mismatch)", "0", "0%", "스크립트 텍스트 오류"],
                     ["Type F (Envelope Violation)", "3", "1.6%", "오디오 envelope 미달"],
                 ])


def section_5_model_training(doc):
    """5. Model Training."""
    add_heading(doc, "5. 모델 학습 (Model Training)", level=1)

    # 5.1 Architecture
    add_heading(doc, "5.1 모델 아키텍처 (Model Architecture)", level=2)
    add_styled_paragraph(
        doc,
        "F5-TTS (Flow-matching 기반 TTS)의 v1 Base 모델을 한국어로 fine-tuning합니다. "
        "F5-TTS는 중국어/영어로 사전학습된 모델이며, 한국어 지원을 위해 vocabulary 확장이 필요합니다."
    )

    arch_items = [
        ("Base Model: ", "F5-TTS v1 Base (flow-matching 기반)"),
        ("사전학습: ", "중국어/영어 음성 데이터 (1,250,000 steps)"),
        ("Vocab 확장: ", "2,545 tokens (원본) → 3,391 tokens (한국어 846자 추가)"),
        ("Embedding 확장: ", "text_embed.weight: 2,546 → 3,392 rows (blank token 포함)"),
        ("Fine-tune 방식: ", "--finetune 플래그, 사전학습 가중치에서 이어서 학습"),
    ]
    for prefix, text in arch_items:
        add_bullet(doc, text, bold_prefix=prefix)

    add_styled_paragraph(doc, "Vocabulary 병합 과정:", bold=True, space_before=Pt(8))
    add_styled_paragraph(
        doc,
        "사전학습 모델의 기존 vocab (2,545 tokens, 중국어 pinyin + 영어)에 "
        "한국어 고유 문자 846자를 추가하여 병합합니다. 기존 vocab을 유지한 채 "
        "신규 토큰만 append하므로 사전학습 임베딩이 보존됩니다. "
        "expand_model_embeddings() 함수로 임베딩 레이어를 자동 확장합니다."
    )

    # 5.2 Infrastructure
    add_heading(doc, "5.2 학습 인프라 (Training Infrastructure)", level=2)

    create_table(doc,
                 ["항목", "사양"],
                 [
                     ["클라우드 플랫폼", "RunPod"],
                     ["GPU", "NVIDIA RTX 4090 (24GB VRAM)"],
                     ["Container Disk", "50 GB"],
                     ["시간당 비용", "$0.59/hr"],
                     ["Framework", "PyTorch 2.0+ / Accelerate (fp16 mixed precision)"],
                     ["총 학습 시간 (Phase 1)", "~2시간"],
                     ["총 비용 (Phase 1)", "~$1.43 (설정+학습+추론 테스트 포함)"],
                 ])

    add_styled_paragraph(doc, "비용 분석:", bold=True, space_before=Pt(8))
    create_table(doc,
                 ["단계", "소요 시간", "비용 (RTX 4090 $0.59/hr)"],
                 [
                     ["환경 설정 + 데이터 다운로드", "~15분", "~$0.15"],
                     ["학습 (100 epochs, batch 6400)", "~2시간", "~$1.18"],
                     ["추론 테스트", "~10분", "~$0.10"],
                     ["합계", "~2.5시간", "~$1.43"],
                 ])

    # 5.3 Hyperparameters
    add_heading(doc, "5.3 하이퍼파라미터 (Hyperparameters)", level=2)

    create_table(doc,
                 ["파라미터", "값", "설명"],
                 [
                     ["Learning Rate", "1e-5", "Fine-tuning용 낮은 학습률"],
                     ["Batch Size (per GPU)", "6,400 frames", "RTX 4090 최적화 (VRAM 82%)"],
                     ["Batch Size Type", "frame", "프레임 기반 동적 배칭"],
                     ["Max Samples per Batch", "64", "배치 내 최대 샘플 수"],
                     ["Grad Accumulation", "2", "효과적 배치 = 12,800 frames"],
                     ["Max Grad Norm", "1.0", "Gradient clipping"],
                     ["Epochs (Phase 1)", "100", "~22,900 updates"],
                     ["Warmup Steps", "500", "Linear warmup"],
                     ["Mixed Precision", "fp16", "메모리 효율 + 속도 향상"],
                     ["Tokenizer", "custom", "병합된 한국어 vocab 사용"],
                     ["Checkpoint 저장", "매 2,000 updates", "last checkpoint: 매 500 updates"],
                     ["Checkpoint 유지", "최근 2개", "디스크 절약 (5.1GB/checkpoint)"],
                 ])

    # 5.4 Training Progress
    add_heading(doc, "5.4 학습 진행 상황 (Training Progress)", level=2)

    add_styled_paragraph(doc, "Phase 1 — 완료:", bold=True, color=COLOR_ACCENT)
    create_table(doc,
                 ["항목", "수치"],
                 [
                     ["Epochs", "100"],
                     ["Total Updates", "~22,900"],
                     ["Updates per Epoch", "229"],
                     ["Time per Epoch", "~73초"],
                     ["총 학습 시간", "~2시간"],
                     ["GPU VRAM 사용량", "~20GB / 24GB (82%)"],
                     ["GPU 활용률", "~94%"],
                 ])

    add_styled_paragraph(doc, "Phase 2 — 계획:", bold=True, color=COLOR_PRIMARY)
    add_styled_paragraph(
        doc,
        "Phase 1의 100 epochs (~22,900 steps)는 기본적인 한국어 음성 생성 능력을 확보하기 위한 "
        "초기 학습입니다. 최종 품질 향상을 위해 Phase 2에서 추가 학습을 계획하고 있습니다."
    )
    phase2_items = [
        ("목표 Steps: ", "~100,000 (Phase 1의 ~4.4배)"),
        ("예상 비용: ", "~$5-6 (RTX 4090 기준)"),
        ("예상 시간: ", "~8-10시간"),
        ("학습률 조정: ", "Phase 2에서 learning rate decay 적용 검토"),
    ]
    for prefix, text in phase2_items:
        add_bullet(doc, text, bold_prefix=prefix)

    # 5.5 Checkpoint Management
    add_heading(doc, "5.5 체크포인트 관리 (Checkpoint Management)", level=2)
    add_styled_paragraph(
        doc,
        "RunPod의 임시 환경 특성상, 체크포인트 유실 방지를 위한 자동 백업 시스템을 구축하였습니다."
    )

    ckpt_items = [
        ("자동 백업: ", "30초 간격으로 새 체크포인트를 감지하여 HuggingFace Hub에 자동 업로드"),
        ("저장 경로: ", "DavidVita/F5TTS-Korean-Training (private repository)"),
        ("백업 대상: ", "체크포인트 (.pt), 학습 로그, TensorBoard 데이터, 오디오 샘플"),
        ("디스크 관리: ", "최근 2개 체크포인트만 유지 (5.1GB/개, 50GB 디스크 제한)"),
        ("재개 지원: ", "중단 시 마지막 체크포인트에서 자동 재개"),
    ]
    for prefix, text in ckpt_items:
        add_bullet(doc, text, bold_prefix=prefix)

    # 5.6 Benchmarks
    add_heading(doc, "5.6 성능 벤치마크 (Performance Benchmarks)", level=2)
    add_styled_paragraph(
        doc,
        "RTX 4090에서 batch size 최적화를 수행하여 처리량을 38% 향상시켰습니다."
    )

    create_table(doc,
                 ["메트릭", "Batch 3200", "Batch 6400", "변화"],
                 [
                     ["Batch size (frames/GPU)", "3,200", "6,400", "2x"],
                     ["Effective batch (w/ accum)", "6,400", "12,800", "2x"],
                     ["Updates per epoch", "488", "229", "-53%"],
                     ["Update speed", "~4.5 updates/sec", "~3.1 updates/sec", "-31%"],
                     ["Time per epoch", "~108초 (1m48s)", "~73초 (1m13s)", "-32%"],
                     ["Throughput (frames/sec)", "~14,400", "~19,840", "+38%"],
                     ["VRAM usage", "12.8GB (52%)", "20.0GB (82%)", "+56%"],
                     ["GPU utilization", "82%", "94%", "+15%"],
                     ["Est. 100 epochs", "~3.0시간", "~2.0시간", "-33%"],
                     ["Est. cost", "~$1.77", "~$1.18", "-$0.59"],
                 ])


def section_6_inference(doc):
    """6. Inference & Deployment."""
    add_heading(doc, "6. 추론 및 배포 (Inference & Deployment)", level=1)

    add_heading(doc, "6.1 추론 방법 (Inference Methods)", level=2)
    add_styled_paragraph(
        doc,
        "학습된 모델은 3가지 방법으로 음성을 생성할 수 있습니다."
    )

    methods = [
        ["CLI (Command Line)", "python scripts/inference.py --text \"텍스트\"", "배치 생성, 스크립팅에 적합"],
        ["Gradio Web UI", "F5-TTS 내장 Web UI", "대화형 테스트, 데모"],
        ["Python API", "f5_tts.infer.infer_cli 모듈", "애플리케이션 통합"],
    ]
    create_table(doc,
                 ["방법", "사용법", "용도"],
                 methods)

    add_styled_paragraph(doc, "주요 추론 옵션:", bold=True, space_before=Pt(8))
    inf_items = [
        ("단일 문장 생성: ", '--text "안녕하세요, 반갑습니다."'),
        ("배치 파일 생성: ", "--text-file scripts/sample_texts.txt"),
        ("사용자 참조 음성: ", "--ref-audio my_voice.wav --ref-text \"참조 텍스트\""),
        ("체크포인트 지정: ", "--checkpoint model_22000_ema.safetensors"),
        ("EMA 비활성화: ", "--no-ema (초기 체크포인트에 권장)"),
    ]
    for prefix, text in inf_items:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading(doc, "6.2 체크포인트 경량화 (Checkpoint Pruning)", level=2)
    add_styled_paragraph(
        doc,
        "학습 체크포인트(.pt)에는 optimizer state, scheduler state 등 추론에 불필요한 "
        "데이터가 포함되어 5.1GB에 달합니다. Pruning을 통해 EMA 가중치만 추출하면 "
        "1.3GB의 .safetensors 파일로 경량화할 수 있습니다."
    )

    create_table(doc,
                 ["항목", "학습 체크포인트 (.pt)", "Pruned (.safetensors)"],
                 [
                     ["파일 크기", "5.1 GB", "1.3 GB"],
                     ["포함 내용", "모델 + optimizer + scheduler + EMA", "EMA 가중치만"],
                     ["용도", "학습 재개", "추론 전용"],
                     ["로딩 속도", "느림", "빠름"],
                 ])

    add_styled_paragraph(doc, "EMA (Exponential Moving Average) 참고:", bold=True, space_before=Pt(8))
    ema_items = [
        ("후기 체크포인트 (>50 epochs): ", "EMA 사용 권장 (기본값)"),
        ("초기 체크포인트 (<20 epochs): ", "EMA 미사용(--no-ema)이 더 나을 수 있음"),
        ("이유: ", "EMA 가중치는 서서히 업데이트되므로 초기에는 사전학습 가중치에 치우침"),
    ]
    for prefix, text in ema_items:
        add_bullet(doc, text, bold_prefix=prefix)


def section_7_key_learnings(doc):
    """7. Key Learnings & Insights."""
    add_heading(doc, "7. 핵심 교훈 및 인사이트 (Key Learnings & Insights)", level=1)

    add_heading(doc, "7.1 한국어 특화 과제 (Korean-specific Challenges)", level=2)

    kr_items = [
        ("Whisper temperature=0 금지: ",
         "temperature=0으로 설정하면 한국어 인식이 치명적으로 저하됩니다. "
         "반드시 기본 temperature fallback을 사용해야 합니다."),
        ("전방 탐색만 사용: ",
         "후방 탐색(backward search)은 한국어 텍스트의 유사한 조사/어미 패턴으로 인해 "
         "중복 매칭을 유발합니다. Forward-only search가 안전합니다."),
        ("탐색 범위 제한: ",
         "100~500줄의 넓은 탐색 범위는 한국어 false-match cascading을 야기합니다. "
         "25줄 forward window가 최적입니다."),
        ("GT-prompting 효과: ",
         "Ground-truth 텍스트를 initial_prompt로 사용하면 R1이 거의 2배 향상됩니다 "
         "(34.7% → 64.3%). 한국어 어휘 인식에 극적인 효과가 있습니다."),
        ("한국어 유사도 기준선: ",
         "무작위 한국어 텍스트 쌍의 baseline 유사도가 0.30~0.40으로 높습니다 "
         "(공통 조사/어미 때문). Threshold 설정 시 반드시 고려해야 합니다."),
    ]
    for prefix, text in kr_items:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading(doc, "7.2 데이터 정제 교훈 (Data Cleansing Lessons)", level=2)

    cleanse_items = [
        ("한 번에 하나만 변경: ",
         "Iteration 4에서 다변수 동시 변경으로 진단이 불가능해졌습니다. "
         "Iteration 4c에서 단일 변수(match confirmation)만 유지하여 해결."),
        ("파라미터 regression 주의: ",
         "Iteration 4a에서 명시적으로 revert한 AUDIO_PAD_MS=100이 Iteration 7에서 다시 발견되었습니다. "
         "파라미터 변경 이력 관리가 중요합니다."),
        ("자동 메트릭의 한계: ",
         "R2/R6 자동 검사가 100% 통과했지만, 수동 청취로 audible bleed가 발견되었습니다. "
         "자동 메트릭만으로는 perceptual 품질을 완전히 보장할 수 없습니다."),
        ("Inline verification은 역효과: ",
         "정렬 중 검증을 수행하면 alignment flow를 방해하여 "
         "수량과 품질 모두 저하됩니다 (Iteration 5에서 -4.2%)."),
        ("메타데이터 보호: ",
         "단일 스크립트 재처리 시 'w' 모드로 파일을 열면 다른 스크립트의 "
         "메타데이터가 파괴됩니다. 항상 기존 데이터를 읽고 병합해야 합니다."),
    ]
    for prefix, text in cleanse_items:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading(doc, "7.3 모델 학습 교훈 (Training Lessons)", level=2)

    train_items = [
        ("exp_name은 사전정의만 가능: ",
         "--exp_name은 F5TTS_v1_Base 같은 아키텍처명만 허용. "
         "사용자 정의 이름은 --dataset_name으로 설정."),
        ("Vocab 병합 필수: ",
         "한국어 전용 vocab만 사용하면 임베딩 크기 불일치 오류 발생. "
         "사전학습 vocab에 한국어를 append하는 방식으로 병합해야 합니다."),
        ("Batch size 최적화: ",
         "3,200 → 6,400 frames로 2배 증가 시 처리량 38% 향상, "
         "학습 시간 33% 단축. VRAM 여유가 있다면 batch size를 최대화해야 합니다."),
        ("체크포인트 관리: ",
         "5.1GB/개 체크포인트는 50GB 디스크를 빠르게 소진합니다. "
         "keep_last_n=2로 제한하고, HuggingFace 자동 백업을 병행합니다."),
        ("CUDA OOM 방지: ",
         "8GB VRAM에서 medium → large 모델 전환 시 반드시 CPU 이동 + gc.collect() 필요."),
    ]
    for prefix, text in train_items:
        add_bullet(doc, text, bold_prefix=prefix)


def section_8_next_steps(doc):
    """8. Next Steps & Roadmap."""
    add_heading(doc, "8. 향후 계획 (Next Steps & Roadmap)", level=1)

    add_heading(doc, "8.1 Phase 2 학습 완료", level=2)
    p2_items = [
        ("목표: ", "~100,000 steps까지 추가 학습 (현재 ~22,900 steps)"),
        ("예상 비용: ", "~$5-6 (RTX 4090 기준, ~8-10시간)"),
        ("기대 효과: ", "더 자연스러운 프로소디(prosody)와 음질 향상"),
        ("Learning Rate: ", "Phase 2에서 cosine/linear decay 검토"),
    ]
    for prefix, text in p2_items:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading(doc, "8.2 모델 평가 및 품질 측정", level=2)
    eval_items = [
        ("객관적 평가: ", "MOS (Mean Opinion Score) 청취 테스트"),
        ("비교 평가: ", "사전학습 모델 대비 한국어 품질 개선도 측정"),
        ("다양성 테스트: ", "다양한 문체(뉴스, 대화, 내레이션 등)에서의 성능 확인"),
        ("Robustness 테스트: ", "학습 데이터에 없는 새로운 어휘/문장에 대한 일반화 성능"),
    ]
    for prefix, text in eval_items:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading(doc, "8.3 프로덕션 배포 고려사항", level=2)
    deploy_items = [
        ("API 서버: ", "FastAPI 기반 REST API 구축 검토"),
        ("실시간 합성: ", "Streaming inference를 위한 최적화"),
        ("모델 경량화: ", "Pruned 체크포인트(1.3GB)로 배포, 추가 양자화 검토"),
        ("모니터링: ", "합성 품질 모니터링 및 피드백 시스템 구축"),
    ]
    for prefix, text in deploy_items:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading(doc, "8.4 데이터셋 확장", level=2)
    data_items = [
        ("Script 5 완성: ", "542~1,018번 라인 추가 녹음 (477줄 추가 가능)"),
        ("데이터 증강: ", "Speed perturbation, pitch shifting 등 augmentation 검토"),
        ("다화자 확장: ", "추가 화자 녹음으로 multi-speaker TTS 가능성 탐색"),
    ]
    for prefix, text in data_items:
        add_bullet(doc, text, bold_prefix=prefix)


def section_9_appendix(doc):
    """9. Appendix."""
    add_heading(doc, "9. 부록 (Appendix)", level=1)

    add_heading(doc, "9.1 환경 구성 (Environment Configuration)", level=2)

    add_styled_paragraph(doc, "로컬 환경 (데이터 정제):", bold=True)
    create_table(doc,
                 ["항목", "사양"],
                 [
                     ["OS", "Windows 10 Pro"],
                     ["Python", "3.11"],
                     ["GPU", "NVIDIA RTX 3060 Ti (8GB VRAM)"],
                     ["CUDA", "Available"],
                     ["파일시스템", "exFAT (G: 드라이브)"],
                     ["저장소", "G:\\Projects\\AI_Research\\TTSDataSetCleanser"],
                 ])

    add_styled_paragraph(doc, "클라우드 환경 (모델 학습):", bold=True, space_before=Pt(8))
    create_table(doc,
                 ["항목", "사양"],
                 [
                     ["플랫폼", "RunPod"],
                     ["GPU", "NVIDIA RTX 4090 (24GB VRAM)"],
                     ["Container Disk", "50 GB"],
                     ["Framework", "PyTorch 2.0+ / Accelerate"],
                     ["Mixed Precision", "fp16"],
                     ["F5-TTS 버전", "v1.1.16"],
                 ])

    add_heading(doc, "9.2 주요 의존성 (Key Dependencies)", level=2)

    add_styled_paragraph(doc, "데이터 정제 파이프라인:", bold=True)
    deps_cleanse = [
        ["openai-whisper", "ASR 전사 및 정렬"],
        ["numpy", "오디오 신호 처리"],
        ["soundfile", "WAV 파일 I/O"],
        ["python-Levenshtein", "텍스트 유사도 계산 (CER)"],
        ["torch + CUDA", "GPU 가속"],
    ]
    create_table(doc,
                 ["패키지", "용도"],
                 deps_cleanse)

    add_styled_paragraph(doc, "모델 학습:", bold=True, space_before=Pt(8))
    deps_train = [
        ["f5-tts", "TTS 모델 학습 및 추론"],
        ["accelerate", "분산/혼합정밀도 학습"],
        ["tensorboard", "학습 모니터링"],
        ["huggingface-hub", "체크포인트 백업 및 데이터셋 관리"],
        ["jamo", "한국어 자모 분해"],
    ]
    create_table(doc,
                 ["패키지", "용도"],
                 deps_train)

    add_heading(doc, "9.3 참고 문서 (Reference Documents)", level=2)

    refs = [
        ["TTS_DATASET_PIPELINE_REQUIREMENTS_v02.md", "데이터 정제 파이프라인 요구사항 명세"],
        ["COMPLETE_GUIDE.md", "F5-TTS 한국어 fine-tuning 재현 가이드"],
        ["TRAINING_LOG.md", "모델 학습 전체 과정 기록"],
        ["training_benchmark.md", "Batch size 최적화 벤치마크"],
        ["iteration_log.md", "데이터 정제 R&D 반복 이력 (7회)"],
        ["evaluation_report.json", "최종 평가 보고서 (R1~R6 메트릭)"],
    ]
    create_table(doc,
                 ["문서", "설명"],
                 refs)

    add_heading(doc, "9.4 리포지토리 구조 (Repository Structure)", level=2)

    add_styled_paragraph(doc, "데이터 정제 (TTSDataSetCleanser):", bold=True)
    add_code_block(doc,
                   "G:\\Projects\\AI_Research\\TTSDataSetCleanser\\\n"
                   "  src/align_and_split.py    # Stage 1-2: 정렬 및 분할\n"
                   "  src/evaluate_dataset.py   # Stage 3-4: 검증 및 평가\n"
                   "  src/pipeline_manager.py   # 워크플로 오케스트레이터\n"
                   "  rawdata/audio/            # 원본 오디오 (19 files, ~7.1GB)\n"
                   "  rawdata/Scripts/          # 원본 스크립트 (5 files)\n"
                   "  datasets/wavs/            # 출력: 분할된 WAV (4,200 files)\n"
                   "  datasets/script.txt       # 출력: 메타데이터\n"
                   "  logs/                     # 평가 보고서, iteration 로그")

    add_styled_paragraph(doc, "모델 학습 (korean-tts-research):", bold=True, space_before=Pt(8))
    add_code_block(doc,
                   "G:\\Git\\korean-tts-research\\\n"
                   "  runpod/*.sh               # RunPod 실행 스크립트\n"
                   "  scripts/inference.py      # 추론 스크립트\n"
                   "  scripts/prune_checkpoint.py  # 체크포인트 경량화\n"
                   "  scripts/prepare_korean_dataset.py  # 데이터 전처리\n"
                   "  COMPLETE_GUIDE.md         # 재현 가이드\n"
                   "  TRAINING_LOG.md           # 학습 과정 기록")


# ══════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════

def main():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Default font style ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_KOREAN
    font.size = FONT_SIZE_BODY
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KOREAN)

    # Set heading styles
    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = FONT_KOREAN
        heading_style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KOREAN)
        heading_style.font.color.rgb = COLOR_HEADING
        if level == 1:
            heading_style.font.size = Pt(18)
        elif level == 2:
            heading_style.font.size = Pt(14)
        else:
            heading_style.font.size = Pt(12)

    # ── Build document ──
    create_cover_page(doc)
    section_1_overview(doc)
    doc.add_page_break()
    section_2_data_collection(doc)
    doc.add_page_break()
    section_3_cleansing_pipeline(doc)
    doc.add_page_break()
    section_4_final_dataset(doc)
    doc.add_page_break()
    section_5_model_training(doc)
    doc.add_page_break()
    section_6_inference(doc)
    doc.add_page_break()
    section_7_key_learnings(doc)
    doc.add_page_break()
    section_8_next_steps(doc)
    doc.add_page_break()
    section_9_appendix(doc)

    # ── Footer with page numbers ──
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Page number field
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)

    # ── Save ──
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Report generated: {OUTPUT}")
    print(f"File size: {OUTPUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
