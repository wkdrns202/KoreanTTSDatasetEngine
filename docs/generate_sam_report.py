"""Generate SAM × TTS Pipeline Cross-Analysis Report (.docx)"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Malgun Gothic'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.25

# Set East Asian font
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Malgun Gothic"/>')
    rpr.append(rFonts)
else:
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Malgun Gothic'
    hrpr = hs.element.get_or_add_rPr()
    hrf = hrpr.find(qn('w:rFonts'))
    if hrf is None:
        hrf = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Malgun Gothic"/>')
        hrpr.append(hrf)
    else:
        hrf.set(qn('w:eastAsia'), 'Malgun Gothic')

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "2F5496")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(9.5)
                r.font.name = 'Malgun Gothic'
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    r.font.name = 'Malgun Gothic'
            if ri % 2 == 1:
                set_cell_shading(cell, "D6E4F0")
    if col_widths:
        for ri, row_obj in enumerate(table.rows):
            for ci, w in enumerate(col_widths):
                row_obj.cells[ci].width = Cm(w)
    return table

def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.2 + level * 0.8)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Malgun Gothic'
        run2 = p.add_run(text)
        run2.font.size = Pt(10)
        run2.font.name = 'Malgun Gothic'
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Malgun Gothic'
    return p

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("SAM × TTS 데이터 정제 파이프라인\n교차 분석 보고서")
run.font.size = Pt(24)
run.bold = True
run.font.name = 'Malgun Gothic'
run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Segment Anything (Meta, 2023)의 방법론과\n한국어 TTS 음성 데이터 자동 정제 파이프라인의 구조적 유사성 분석")
run.font.size = Pt(13)
run.font.name = 'Malgun Gothic'
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("2026년 3월 24일\n천영재 (테크 7기 박사과정)")
run.font.size = Pt(11)
run.font.name = 'Malgun Gothic'

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading('목차 (Table of Contents)', level=1)
toc_items = [
    "1. 분석 배경 및 목적",
    "2. 핵심 구조 대비: 구조적 동형성",
    "3. Data Engine — 가장 깊은 유사성",
    "4. 모델 아키텍처 대비",
    "5. 모호성(Ambiguity) 처리 전략 비교",
    "6. Zero-shot Transfer와 재현성 검증",
    "7. 종합 장단점 비교",
    "8. SAM에서 차용 가능한 아이디어 (우선순위순)",
    "9. 결론 및 시사점",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 1. 분석 배경 및 목적
# ============================================================
doc.add_heading('1. 분석 배경 및 목적', level=1)

doc.add_paragraph(
    'Meta AI Research에서 발표한 Segment Anything Model(SAM) 논문(Kirillov et al., 2023)은 '
    '이미지 세그멘테이션 분야의 Foundation Model을 제안하며, promptable segmentation이라는 '
    '새로운 태스크 정의, 반복적 Data Engine을 통한 대규모 데이터셋 구축, 그리고 zero-shot '
    'transfer 능력을 핵심 기여로 제시하였다.'
)
doc.add_paragraph(
    '본 보고서는 SAM의 핵심 방법론과 본 프로젝트(한국어 TTS 음성 데이터 자동 정제 파이프라인)의 '
    '작업 흐름을 교차 분석하여, 두 연구 간의 구조적 유사성과 차이점을 체계적으로 정리하고, '
    'SAM으로부터 도출 가능한 인사이트 및 향후 적용 방안을 제안하는 것을 목적으로 한다.'
)

doc.add_heading('분석 대상', level=3)
add_bullet(doc, 'Kirillov, A. et al. (2023). Segment Anything. Meta AI Research, FAIR.')
add_bullet(doc, '천영재 (2026). 인공지능 음성합성(TTS) 파이프라인 구축: 데이터 수집 및 라벨링 자동화 프로세스 연구개발 요약본.')

# ============================================================
# 2. 핵심 구조 대비
# ============================================================
doc.add_heading('2. 핵심 구조 대비: 구조적 동형성', level=1)

doc.add_paragraph(
    '두 프로젝트는 도메인(이미지 vs. 음성)은 다르지만, 본질적으로 "연속 신호에서 의미 단위를 '
    '정확히 잘라내는" 세그멘테이션 문제이며, Foundation Model을 프롬프트로 구동하고, '
    '반복적 Data Engine으로 품질을 끌어올리는 동일한 패러다임을 공유한다.'
)

add_table(doc,
    ["구조적 요소", "SAM (이미지 세그멘테이션)", "본 프로젝트 (음성 세그멘테이션)"],
    [
        ["과제 정의", "Promptable Segmentation\n— 프롬프트 → 유효한 마스크 반환", "Promptable Alignment\n— 대본(프롬프트) → 해당 음성 구간 반환"],
        ["입력 데이터", "고해상도 이미지 (11M장)", "장시간 고해상도 오디오\n(5개 대본, 수천 문장)"],
        ["\"세그먼트\"의 의미", "이미지 내 객체 영역 (마스크)", "오디오 내 문장 구간 (WAV 클립)"],
        ["Foundation Model", "ViT-H (MAE 사전학습)\n→ Image Encoder", "Whisper medium/large\n→ ASR Encoder"],
        ["프롬프트 체계", "점, 박스, 마스크, 텍스트", "대본 텍스트\n(유사도 기반 매칭 프롬프트)"],
        ["품질 보증", "IoU 기반 confidence score\n+ NMS 필터링", "6개 요구사항 (R1~R6)\n+ 다단계 검증"],
        ["반복 개선", "Data Engine 3단계\n(수동→반자동→전자동)", "6회 Iteration\n(35% → 95%+)"],
        ["최종 산출물", "SA-1B 데이터셋\n(1.1B 마스크, 11M 이미지)", "4,200문장 음성-텍스트 쌍\n데이터셋 (8시간)"],
    ],
    col_widths=[3.5, 6.0, 6.0]
)

# ============================================================
# 3. Data Engine
# ============================================================
doc.add_heading('3. Data Engine — 가장 깊은 유사성', level=1)

doc.add_paragraph(
    'SAM 논문의 가장 핵심적인 기여는 Data Engine 개념이다. 모델이 어노테이션을 보조하고, '
    '새로운 데이터가 모델을 개선하는 선순환 구조로, 본 프로젝트의 6-Stage 반복 파이프라인과 '
    '구조적으로 거의 동일하다.'
)

doc.add_heading('3.1 SAM Data Engine 3단계', level=2)
add_bullet(doc, ' SAM이 보조하며 사람이 마스크를 직접 생성 → 4.3M 마스크, 모델 6회 재학습', 'Manual (수동):')
add_bullet(doc, ' SAM이 먼저 마스크를 예측하고, 사람은 누락분만 보완 → 5.9M 마스크 추가', 'Semi-automatic (반자동):')
add_bullet(doc, ' 32×32 그리드 포인트로 전체 자동 생성 → 1.1B 마스크', 'Fully automatic (전자동):')

doc.add_heading('3.2 본 프로젝트 Iteration 흐름', level=2)
add_bullet(doc, ' 기본 Whisper 전사 + 단순 매칭 → 정확도 35% (≈ SAM Manual 단계)', '1~2차 (초기):')
add_bullet(doc, ' 매칭 임계값 조정, 프롬프팅, 대형 모델 2차 검증 도입 → 64~72% (≈ Semi-automatic)', '3~5차 (중간 개선):')
add_bullet(doc, ' 경계 최적화, envelope 표준화, 전체 자동 파이프라인 확립 → 95%+ (≈ Fully automatic)', '6차 (최종):')

doc.add_heading('3.3 Data Engine 세부 대비', level=2)
add_table(doc,
    ["대비 항목", "SAM", "본 프로젝트"],
    [
        ["모델 재학습/개선 횟수", "6회 (모델 retrain)", "6회 (알고리즘 iteration)"],
        ["자동화 전환", "점진적 (수동→자동)", "점진적 (수동 튜닝→자동 파이프라인)"],
        ["피드백 루프", "새 데이터 → 모델 재학습\n→ 더 좋은 데이터", "실패 분석(Stage 6)\n→ 알고리즘 개선(Stage 1)\n→ 더 좋은 정렬"],
        ["품질 필터링", "IoU threshold (88.0)\n+ stability check", "similarity threshold (0.50)\n+ Tier 2 대형 모델 검증"],
        ["어노테이션 시간 변화", "34초/마스크 → 14초/마스크", "수작업 100% → 자동 95%+"],
    ],
    col_widths=[4.0, 5.5, 5.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('핵심 차이: ')
run.bold = True
run.font.name = 'Malgun Gothic'
p.add_run(
    'SAM은 모델 자체를 재학습(retrain)하지만, 본 프로젝트는 모델(Whisper)은 고정하고 '
    '알고리즘/파라미터를 개선한다. 이는 리소스 차이(Meta AI vs 개인 연구)에서 오는 합리적 '
    '선택이며, 동시에 "모델 고정 하에서 알고리즘 최적화만으로 Data Engine을 구동할 수 있다"는 '
    '독자적 기여를 보여준다.'
).font.name = 'Malgun Gothic'

# ============================================================
# 4. 모델 아키텍처 대비
# ============================================================
doc.add_heading('4. 모델 아키텍처 대비', level=1)

doc.add_heading('4.1 SAM 3-컴포넌트 구조', level=2)
add_bullet(doc, ' ViT-H 기반, 무거움, 이미지당 1회 실행 → Image Embedding 생성', 'Image Encoder:')
add_bullet(doc, ' 경량, 점/박스/텍스트를 256차원 벡터로 변환', 'Prompt Encoder:')
add_bullet(doc, ' 경량 (50ms), 임베딩 조합 → 세그멘테이션 마스크 × 3 출력 (모호성 대응)', 'Mask Decoder:')

doc.add_heading('4.2 본 프로젝트의 대응 구조', level=2)
add_bullet(doc, ' Whisper medium, 무거움, 오디오 파일당 1회 전사 → ASR 세그먼트 + 타임스탬프', 'Audio Encoder:')
add_bullet(doc, ' 경량, 텍스트 유사도 계산으로 대본-세그먼트 매칭 (프롬프트 역할)', 'Script Matcher:')
add_bullet(doc, ' 경량, 매칭된 타임스탬프 기반 오디오 분할 + 경계 처리 → WAV 클립 출력', 'Audio Splitter:')

doc.add_heading('4.3 구조적 동형성', level=2)
doc.add_paragraph(
    '두 시스템 모두 "무거운 인코더가 한 번 실행되고, 경량 디코더/매처가 여러 프롬프트에 대해 '
    '빠르게 동작하는" amortized computation 패턴을 공유한다. SAM에서 이미지 임베딩을 한 번 '
    '계산하고 여러 프롬프트를 시도하는 것처럼, 본 프로젝트에서도 Whisper 전사를 한 번 수행하고 '
    '대본의 각 라인을 프롬프트로 순차 매칭한다.'
)

add_table(doc,
    ["비교 요소", "SAM", "본 프로젝트"],
    [
        ["무거운 연산 (1회)", "ViT-H Image Encoding", "Whisper ASR 전사"],
        ["경량 연산 (N회)", "Prompt Encoding + Mask Decoding\n(50ms/프롬프트)", "텍스트 유사도 매칭\n+ 오디오 분할"],
        ["Amortized 효과", "동일 이미지에 다양한 프롬프트 적용", "동일 전사 결과에 대본 전체 매칭"],
    ],
    col_widths=[4.0, 5.5, 5.5]
)

# ============================================================
# 5. 모호성 처리
# ============================================================
doc.add_heading('5. 모호성(Ambiguity) 처리 전략 비교', level=1)

doc.add_heading('5.1 SAM의 접근', level=2)
add_bullet(doc, '하나의 점 프롬프트가 여러 객체(예: 셔츠 vs 사람 전체)를 지칭할 수 있음')
add_bullet(doc, '3개의 유효 마스크를 동시 출력 (whole / part / subpart)')
add_bullet(doc, 'IoU confidence score로 순위 매김')
add_bullet(doc, '"모호해도 최소 하나는 유효한 답"을 보장하는 설계 철학')

doc.add_heading('5.2 본 프로젝트의 현재 상황', level=2)
add_bullet(doc, '한국어 조사/어미 유사성으로 오탐(false positive) 발생 가능')
add_bullet(doc, '하나의 대본 라인이 여러 ASR 세그먼트에 걸칠 수 있음 (세그먼트 병합 필요)')
add_bullet(doc, '현재는 단일 최선 매칭만 선택 (MAX_MERGE=5로 병합 시도)')

doc.add_heading('5.3 SAM에서 도출한 시사점', level=2)

p = doc.add_paragraph()
run = p.add_run('a) Multi-hypothesis 매칭: ')
run.bold = True
p.add_run(
    'Greedy 최선 하나 대신, SAM처럼 상위 N개 후보 매칭을 유지하고, '
    '후속 문맥(다음 대본 라인의 매칭 결과)까지 본 뒤 최종 선택. '
    '특히 re-sync 실패를 줄일 수 있음.'
)

p = doc.add_paragraph()
run = p.add_run('b) Confidence Score 고도화: ')
run.bold = True
p.add_run(
    'SAM은 각 마스크에 예측 IoU를 부여한다. 본 프로젝트의 similarity score도 유사한 역할이나, '
    '"이 매칭이 맞을 확률" 자체를 별도로 추정하는 메커니즘을 도입하면 임계값 튜닝 의존도를 줄일 수 있음.'
)

p = doc.add_paragraph()
run = p.add_run('c) Stability Check: ')
run.bold = True
p.add_run(
    'SAM은 threshold를 ±δ 변화시켜도 마스크가 안정적인지 확인한다. '
    '본 프로젝트에서도 similarity threshold를 ±0.05 변동시켜 매칭 결과가 바뀌지 않는 '
    '"안정 매칭"만 자동 확정하고, 불안정한 것만 Tier 2 검증에 넘기는 전략이 가능하다.'
)

# ============================================================
# 6. Zero-shot Transfer
# ============================================================
doc.add_heading('6. Zero-shot Transfer와 재현성 검증', level=1)

doc.add_paragraph(
    'SAM의 핵심 가치 중 하나는 학습에 사용하지 않은 23개 데이터셋에서도 유효하게 작동한다는 '
    'zero-shot transfer 능력이다. 본 프로젝트의 로드맵에도 "정제 모델 재현성 검증 — 다른 '
    '데이터셋에도 일관적으로 적용 가능한지"가 단기 목표로 명시되어 있으며, 이는 SAM의 '
    'zero-shot transfer 개념에 정확히 대응한다.'
)

doc.add_heading('6.1 SAM이 zero-shot을 달성한 핵심 요인', level=2)
add_bullet(doc, ' "프롬프트 → 유효 마스크"라는 범용적 태스크 정의', 'Task 정의의 일반성:')
add_bullet(doc, ' 11M 이미지, 다양한 도메인', '대규모 다양한 데이터:')
add_bullet(doc, ' 다운스트림 태스크를 프롬프트로 변환', '프롬프트 엔지니어링:')

doc.add_heading('6.2 본 프로젝트의 재현성 확보를 위한 시사점', level=2)
add_bullet(doc, '현재 파이프라인은 한국어 특화 파라미터(조사 유사성 대응, ASR 디코딩 전략)에 의존')
add_bullet(doc, '다른 언어/데이터셋으로 전이하려면, SAM처럼 언어에 독립적인 태스크 정의가 필요')
add_bullet(doc, '"ASR 전사 → 텍스트 유사도 매칭 → 세그먼트 추출"이라는 범용 프레임워크와, 언어별 프롬프트/파라미터를 플러그인으로 분리하는 설계가 핵심')

# ============================================================
# 7. 종합 장단점 비교
# ============================================================
doc.add_heading('7. 종합 장단점 비교', level=1)

doc.add_heading('7.1 각 방식의 장점', level=2)
add_table(doc,
    ["측면", "SAM 방식의 장점", "본 프로젝트 방식의 장점"],
    [
        ["확장성", "모델 재학습으로\n데이터↔모델 시너지", "모델 고정으로\n재현성 높음, 리소스 절약"],
        ["모호성 대응", "Multi-output으로 근본적 해결", "단일 출력이지만\n도메인 지식으로 보완"],
        ["품질 보증", "IoU + 안정성\n자동 필터링", "6개 요구사항 체계\n+ Tier 2 대형 모델 검증"],
        ["비용", "대규모 인프라 활용", "단일 RTX 3060 Ti로\n수행 가능"],
        ["도메인 전이", "23개 데이터셋\nzero-shot 입증", "한국어 특화 최적화로\n해당 도메인 최고 성능"],
        ["자동화 수준", "최종 100% 자동", "95% 자동 + 5% 수동 검수"],
    ],
    col_widths=[3.0, 5.5, 5.5]
)

doc.add_heading('7.2 각 방식의 한계', level=2)
add_table(doc,
    ["측면", "SAM 방식의 한계", "본 프로젝트 방식의 한계"],
    [
        ["리소스", "256 GPU × 대규모 학습 필요", "8GB VRAM 제약\n(Tier 2 전환 시 OOM)"],
        ["도메인 깊이", "범용이지만 특정 도메인에서\n전문 모델에 열세", "한국어 TTS에 최적화,\n다른 언어 전이 미검증"],
        ["잔여 오류", "미세 객체/모호한 경계에서 한계", "ASR의 한국어 전문 어휘\n인식 한계"],
    ],
    col_widths=[3.0, 5.5, 5.5]
)

# ============================================================
# 8. 차용 가능한 아이디어
# ============================================================
doc.add_heading('8. SAM에서 차용 가능한 아이디어 (우선순위순)', level=1)

doc.add_heading('8.1 [논문 작성] Data Engine 3단계 프레임워크 공식화', level=2)
doc.add_paragraph(
    '본 프로젝트의 6회 iteration을 SAM의 3단계(Manual → Semi-auto → Fully-auto)로 '
    '재구조화하여 서술하면 연구 기여의 명확성이 크게 향상된다.'
)
add_bullet(doc, ' Iteration 1~2 (기본 ASR + 수동 파라미터 튜닝)', 'Manual:')
add_bullet(doc, ' Iteration 3~5 (Tier 2 모델 도입, 자동 재동기화)', 'Semi-automatic:')
add_bullet(doc, ' Iteration 6 (전체 자동 파이프라인 + 자동 품질 게이팅)', 'Fully automatic:')

doc.add_heading('8.2 [구현] Stability-based Filtering 도입', level=2)
doc.add_paragraph(
    'SAM의 mask stability check를 음성 정렬에 적용하는 방안이다.'
)
add_bullet(doc, '매칭 결과가 threshold ± δ에서도 동일 → "안정 매칭" (자동 확정)')
add_bullet(doc, '매칭 결과가 threshold 변동에 민감 → "불안정 매칭" (Tier 2 검증 대상)')
doc.add_paragraph(
    '이를 통해 Tier 2(large 모델)을 전체가 아닌 불안정 매칭에만 적용하여, 8GB VRAM 환경에서의 '
    '비용 효율성을 더욱 높일 수 있다.'
)

doc.add_heading('8.3 [중기 개선] Ambiguity-aware Multi-output', level=2)
doc.add_paragraph(
    '매칭 후보를 1개가 아닌 3개 유지하고, 전후 문맥 일관성(monotonic alignment 제약)으로 '
    '최종 선택하는 전략. 특히 re-sync 상황에서 효과적이며, 연속 매칭 실패(CONSEC_FAIL_LIMIT) '
    '발생 빈도를 줄일 수 있다.'
)

doc.add_heading('8.4 [장기 비전] Compositionality — 모듈식 컴포넌트화', level=2)
doc.add_paragraph(
    'SAM이 DALL·E 등 다른 시스템의 컴포넌트로 활용되듯, 본 프로젝트의 정제 파이프라인도 '
    'TTS 훈련의 모듈식 전처리 컴포넌트로 패키징하는 것이 장기 가치이다. SAM 논문 Section 8에서 '
    'compositionality를 foundation model의 핵심 가치로 강조하고 있으며, 이는 본 프로젝트의 '
    '"범용 프레임워크 확립" 목표와 정확히 일치한다.'
)

# ============================================================
# 9. 결론
# ============================================================
doc.add_heading('9. 결론 및 시사점', level=1)

doc.add_paragraph(
    'SAM 논문은 이미지 도메인이지만, 본 프로젝트와 높은 구조적 동형성을 보인다.'
)

p = doc.add_paragraph()
p.style = 'Intense Quote' if 'Intense Quote' in [s.name for s in doc.styles] else 'Normal'
run = p.add_run(
    '"연속 신호 → Foundation Model로 전사/인코딩 → 프롬프트 기반 세그멘테이션 '
    '→ 반복적 Data Engine으로 품질 향상"'
)
run.italic = True
run.font.size = Pt(11)
run.font.name = 'Malgun Gothic'

doc.add_paragraph(
    '이 패러다임이 비전에서 입증된 것을 음성 도메인에 독립적으로 재발견/적용했다는 점이 '
    '본 프로젝트의 연구적 기여를 강화하는 근거가 된다.'
)

doc.add_paragraph(
    '논문 작성 시 SAM을 cross-domain reference로 인용하면서, 본 프로젝트가 음성 도메인에서 '
    '유사한 Data Engine 패러다임을 한국어 특화 과제에 맞게 변형 적용했음을 주장할 수 있다.'
)

p = doc.add_paragraph()
run = p.add_run('독자성의 근거: ')
run.bold = True
run.font.name = 'Malgun Gothic'
p.add_run(
    'SAM은 모델을 재학습하는 방식이고, 본 프로젝트는 모델을 고정하고 알고리즘을 개선하는 '
    '방식이라는 차이가 있어, 단순 모방이 아닌 도메인 적응적 변형으로서의 독자성을 보여준다. '
    '제한된 리소스(단일 GPU) 환경에서 Foundation Model의 재학습 없이 Data Engine 패러다임을 '
    '성공적으로 구동한 사례로서 의미가 있다.'
).font.name = 'Malgun Gothic'

# ============================================================
# References
# ============================================================
doc.add_heading('참고문헌', level=1)
refs = [
    'Kirillov, A., Mintun, E., Ravi, N., Mao, H., Rolland, C., Gustafson, L., Xiao, T., '
    'Whitehead, S., Berg, A.C., Lo, W.-Y., Dollár, P., & Girshick, R. (2023). '
    'Segment Anything. Meta AI Research, FAIR.',

    '천영재 (2026). 인공지능 음성합성(TTS) 파이프라인 구축: 데이터 수집 및 라벨링 자동화 '
    '프로세스 연구개발 요약본. 연구 진행 보고서.',

    'Radford, A. et al. (2023). Robust Speech Recognition via Large-Scale Weak Supervision. '
    'Proc. ICML 2023. (Whisper)',

    'Bommasani, R. et al. (2021). On the Opportunities and Risks of Foundation Models. '
    'arXiv:2108.07258.',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(f'[{i}] {ref}')
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(9.5)

# -- Save --
output_path = os.path.join(os.path.dirname(__file__), 'SAM_TTS_Cross_Analysis_Report.docx')
doc.save(output_path)
print(f'Report saved to: {output_path}')
