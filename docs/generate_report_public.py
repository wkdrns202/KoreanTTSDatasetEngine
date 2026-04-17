"""
Korean TTS Research Progress Report — PUBLIC Version
Abstracted for external distribution. No proprietary parameters, hardware details, or R&D know-how exposed.
"""

from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Paths ──────────────────────────────────────────────────────────────
BASE = Path(__file__).resolve().parent.parent
OUTPUT = BASE / "docs" / "TTS_Research_Progress_Report_Public.docx"

# ── Constants ──────────────────────────────────────────────────────────
FONT_KOREAN = "Malgun Gothic"
FONT_CODE = "Consolas"
FONT_SIZE_BODY = Pt(10.5)
FONT_SIZE_SMALL = Pt(9)
FONT_SIZE_TABLE = Pt(9.5)

COLOR_PRIMARY = RGBColor(0x1A, 0x56, 0xDB)
COLOR_HEADING = RGBColor(0x1F, 0x29, 0x37)
COLOR_ACCENT = RGBColor(0x2E, 0x7D, 0x32)
COLOR_GRAY = RGBColor(0x66, 0x66, 0x66)
COLOR_TABLE_HEADER = "1A56DB"
COLOR_TABLE_ALT = "F0F4FA"
COLOR_PASS = RGBColor(0x2E, 0x7D, 0x32)


# ══════════════════════════════════════════════════════════════════════
#  UTILITIES (same as internal version)
# ══════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, font_name=FONT_KOREAN, size=FONT_SIZE_BODY, bold=False, color=None, italic=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_styled_paragraph(doc, text, font_name=FONT_KOREAN, size=FONT_SIZE_BODY,
                         bold=False, color=None, alignment=None, space_after=Pt(6),
                         space_before=Pt(0), italic=False):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold, color, italic)
    return p


def add_bullet(doc, text, level=0, bold_prefix="", font_size=FONT_SIZE_BODY):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_run_font(run, size=font_size, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=font_size)
    return p


def create_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=FONT_SIZE_TABLE, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, COLOR_TABLE_HEADER)
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_run_font(run, size=FONT_SIZE_TABLE)
            if r_idx % 2 == 1:
                set_cell_shading(cell, COLOR_TABLE_ALT)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_KOREAN
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KOREAN)
    return h


# ══════════════════════════════════════════════════════════════════════
#  DOCUMENT SECTIONS — PUBLIC VERSION
# ══════════════════════════════════════════════════════════════════════

def create_cover_page(doc):
    for _ in range(6):
        doc.add_paragraph()

    add_styled_paragraph(
        doc, "한국어 TTS 연구 진행 보고서",
        size=Pt(28), bold=True, color=COLOR_HEADING,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12)
    )
    add_styled_paragraph(
        doc, "Korean Text-to-Speech Research Progress Report",
        size=Pt(16), color=COLOR_PRIMARY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8)
    )
    add_styled_paragraph(
        doc, "— Public Summary —",
        size=Pt(12), color=COLOR_GRAY, italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(48)
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    set_run_font(run, size=Pt(12), color=COLOR_PRIMARY)
    doc.add_paragraph()

    meta_items = [
        ("프로젝트", "한국어 TTS 음성합성 파이프라인 구축"),
        ("작성일", datetime.now().strftime("%Y년 %m월 %d일")),
        ("문서 분류", "공개용 요약 (Public Summary)"),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{label}: ")
        set_run_font(run, size=Pt(12), bold=True, color=COLOR_HEADING)
        run = p.add_run(value)
        set_run_font(run, size=Pt(12), color=COLOR_GRAY)

    doc.add_page_break()


def section_1_overview(doc):
    add_heading(doc, "1. 프로젝트 개요 (Project Overview)", level=1)

    add_heading(doc, "1.1 목적 및 배경 (Purpose & Background)", level=2)
    add_styled_paragraph(
        doc,
        "본 프로젝트는 고품질 한국어 TTS(Text-to-Speech) 음성합성 시스템 구축을 목표로 합니다. "
        "전문 성우의 녹음 데이터를 기반으로, 자동화된 데이터 정제 파이프라인을 통해 학습용 데이터셋을 "
        "생성하고, 최신 TTS 모델을 fine-tuning하여 자연스러운 한국어 음성을 합성합니다."
    )
    add_styled_paragraph(
        doc,
        "기존 한국어 TTS 데이터셋 구축은 수작업 레이블링에 의존하여 비용과 시간이 많이 소요되었습니다. "
        "본 연구에서는 ASR(Automatic Speech Recognition) 기반 자동 정렬 기법을 통해 "
        "이 과정을 자동화하고, 반복적 R&D 사이클을 통해 높은 정확도를 달성하였습니다."
    )

    add_heading(doc, "1.2 전체 파이프라인 구조 (End-to-End Pipeline)", level=2)

    pipeline_flow = (
        "녹음 (Record)  →  데이터 정제 (Cleanse)  →  모델 학습 (Train)  →  추론 (Inference)"
    )
    add_styled_paragraph(
        doc, pipeline_flow,
        size=Pt(12), bold=True, color=COLOR_PRIMARY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(12), space_after=Pt(12)
    )

    steps = [
        ("1단계 - 데이터 수집: ", "전문 성우 녹음, 다수의 장시간 오디오 파일 및 대본 확보"),
        ("2단계 - 데이터 정제: ", "ASR 기반 자동 정렬, 다단계 파이프라인, 반복적 품질 개선"),
        ("3단계 - 모델 학습: ", "최신 TTS 모델 fine-tuning, 클라우드 GPU 활용"),
        ("4단계 - 추론/배포: ", "다양한 인터페이스(CLI, Web UI, API) 지원"),
    ]
    for prefix, text in steps:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading(doc, "1.3 기술 스택 요약 (Technology Stack)", level=2)

    create_table(doc,
                 ["분류", "접근 방식"],
                 [
                     ["음성 인식 (ASR)", "최신 multilingual ASR 모델 활용"],
                     ["TTS 모델", "Flow-matching 기반 최신 TTS 아키텍처"],
                     ["학습 인프라", "클라우드 GPU (on-demand)"],
                     ["오디오 처리", "고해상도 오디오 (48kHz 이상, 무손실)"],
                     ["모델 관리", "HuggingFace Hub 기반 버전 관리 및 배포"],
                 ])


def section_2_data_collection(doc):
    add_heading(doc, "2. 데이터 수집 (Data Collection)", level=1)

    add_heading(doc, "2.1 데이터 개요 (Data Overview)", level=2)
    add_styled_paragraph(
        doc,
        "전문 성우 1인의 고품질 스튜디오 녹음을 기반으로 데이터를 구축하였습니다. "
        "다양한 문체(문학, 교육, 내러티브 등)를 포괄하는 대본을 사용하여 "
        "모델의 일반화 성능을 확보하고자 하였습니다."
    )

    create_table(doc,
                 ["항목", "사양"],
                 [
                     ["녹음 규모", "수천 문장 규모 (다수의 장시간 오디오)"],
                     ["오디오 품질", "스튜디오 녹음, 고해상도 무손실 포맷"],
                     ["채널", "Mono"],
                     ["녹음자", "전문 성우 1인"],
                 ])

    add_heading(doc, "2.2 대본 구성 (Script Composition)", level=2)
    add_styled_paragraph(
        doc,
        "5개의 주제별 대본으로 구성되며, 각 대본은 수백~천 단위의 한국어 문장을 포함합니다. "
        "다양한 도메인을 포괄하여 TTS 모델이 특정 문체에 편향되지 않도록 설계하였습니다."
    )

    create_table(doc,
                 ["대본", "내용 특성"],
                 [
                     ["대본 1", "문학 작품 낭독"],
                     ["대본 2", "IT/기술 교육 콘텐츠"],
                     ["대본 3", "고전 및 현대 문학 텍스트"],
                     ["대본 4", "일반 교양 콘텐츠"],
                     ["대본 5", "내러티브/서사 콘텐츠"],
                 ])

    add_styled_paragraph(
        doc,
        "다양한 문체를 포함함으로써, 뉴스 읽기, 대화체, 감정 표현 등 "
        "실제 활용 시나리오에 폭넓게 대응할 수 있는 데이터셋을 목표로 하였습니다.",
        space_before=Pt(4)
    )


def section_3_cleansing_pipeline(doc):
    add_heading(doc, "3. 데이터 정제 파이프라인 (Data Cleansing Pipeline)", level=1)

    add_styled_paragraph(
        doc,
        "데이터 정제는 본 프로젝트의 핵심 연구 영역입니다. 자동화된 다단계 파이프라인을 설계하여 "
        "정렬, 검증, 진단, 개선의 반복 사이클을 구현하였으며, 다수의 반복(iteration)을 통해 "
        "95%를 초과하는 정렬 정확도를 달성하였습니다."
    )

    # 3.1 Architecture
    add_heading(doc, "3.1 파이프라인 아키텍처 (Pipeline Architecture)", level=2)
    add_styled_paragraph(
        doc,
        "파이프라인은 6개의 Stage로 구성된 반복형(iterative) 구조를 채택하였습니다. "
        "품질 기준 미달 시 자동으로 진단 및 개선 사이클을 수행합니다."
    )

    create_table(doc,
                 ["Stage", "기능", "설명"],
                 [
                     ["Stage 1", "Align & Split", "ASR로 장시간 음성을 전사한 후, 대본과 자동 정렬하여 문장 단위로 분할"],
                     ["Stage 2", "Post-process", "분할된 오디오의 경계 처리, 무음 구간 표준화, 음질 보정"],
                     ["Stage 3", "Validate", "분할 결과를 독립적으로 재검증하여 정렬 품질 측정"],
                     ["Stage 4", "Evaluate", "전체 품질 메트릭 집계 및 요구사항 충족 여부 판정"],
                     ["Stage 5", "Finalize", "품질 기준 충족 시 최종 데이터셋 확정"],
                     ["Stage 6", "Diagnose & Improve", "미충족 시 실패 원인 분석 후 파이프라인 개선"],
                 ])

    add_styled_paragraph(
        doc,
        "Stage 6에서 Stage 1으로 되돌아가는 피드백 루프를 통해, 각 iteration마다 "
        "실패 패턴을 분석하고 알고리즘을 점진적으로 개선합니다."
    )

    # 3.2 Core Methodology
    add_heading(doc, "3.2 핵심 방법론 (Core Methodology)", level=2)
    add_styled_paragraph(
        doc,
        "장시간 연속 녹음 오디오에서 개별 문장을 정확하게 추출하는 것이 핵심 과제입니다. "
        "ASR 모델의 전사 결과와 원본 대본 텍스트를 자동으로 정렬(alignment)하는 "
        "독자적인 알고리즘을 개발하였습니다."
    )

    add_styled_paragraph(doc, "주요 기술적 접근:", bold=True, space_before=Pt(8))

    methods = [
        ("ASR 기반 자동 정렬: ",
         "최신 multilingual ASR 모델을 활용하여 장시간 오디오를 전사하고, "
         "원본 대본과의 텍스트 유사도 기반 매칭을 수행합니다."),
        ("한국어 특화 탐색 전략: ",
         "한국어 텍스트의 특성(조사, 어미 유사성)을 고려한 방향성 탐색 알고리즘을 적용하여 "
         "오정렬(misalignment)을 방지합니다."),
        ("세그먼트 병합 최적화: ",
         "ASR 모델이 하나의 문장을 여러 세그먼트로 분할하는 문제를 해결하기 위해 "
         "연속 세그먼트 결합 매칭을 수행합니다."),
        ("자동 재동기화: ",
         "연속적인 매칭 실패 발생 시 자동으로 탐색 범위를 확대하여 "
         "정렬 복구를 시도하는 안전장치를 구현하였습니다."),
        ("다단계 검증: ",
         "1차 모델과 2차 대형 모델의 계층적 검증을 통해 "
         "단일 모델의 인식 한계를 보완합니다."),
    ]
    for prefix, text in methods:
        add_bullet(doc, text, bold_prefix=prefix)

    # 3.3 Audio Post-processing
    add_heading(doc, "3.3 오디오 후처리 (Audio Post-processing)", level=2)
    add_styled_paragraph(
        doc,
        "TTS 학습에 적합한 오디오 품질을 확보하기 위해, 분할된 각 WAV 파일에 대해 "
        "체계적인 후처리를 적용합니다."
    )

    postproc = [
        ("경계 처리: ", "절단점 최적화 및 fade 처리로 클릭/팝 아티팩트 제거"),
        ("무음 구간 표준화: ", "TTS 학습에 최적화된 일관된 오디오 envelope 적용 "
         "(음성 전후 표준 무음 구간 보장)"),
        ("음량 정규화: ", "음성 구간의 peak 음량을 표준화하여 학습 데이터 일관성 확보"),
        ("품질 보존: ", "원본 오디오의 고해상도 포맷을 유지한 채 처리"),
    ]
    for prefix, text in postproc:
        add_bullet(doc, text, bold_prefix=prefix)

    # 3.4 Quality Assurance
    add_heading(doc, "3.4 품질 보증 체계 (Quality Assurance)", level=2)
    add_styled_paragraph(
        doc,
        "6가지 품질 요구사항(R1~R6)을 정의하고, 모든 항목에서 95% 이상을 달성해야 "
        "데이터셋으로 확정되는 엄격한 품질 관리 체계를 운용합니다."
    )

    create_table(doc,
                 ["요구사항", "설명"],
                 [
                     ["R1: 정렬 정확도", "각 오디오 파일의 내용이 대본 텍스트와 정확히 일치하는지 검증"],
                     ["R2: 경계 품질", "오디오 경계에 인접 문장의 음성이나 잡음이 없는지 검증"],
                     ["R3: 종합 통과율", "R1, R2, R6를 모두 동시에 충족하는 세그먼트의 비율"],
                     ["R4: 메타데이터 무결성", "모든 오디오 파일에 정확한 메타데이터가 존재하는지 검증"],
                     ["R5: 재현성", "동일 입력에 대해 동일 결과를 보장"],
                     ["R6: 오디오 규격", "TTS 학습에 필요한 표준 오디오 envelope 충족 여부"],
                 ])

    # 3.5 R&D Iteration Summary
    add_heading(doc, "3.5 반복적 품질 개선 (Iterative Improvement)", level=2)
    add_styled_paragraph(
        doc,
        "초기 파이프라인의 정렬 정확도는 35% 수준이었으나, "
        "반복적인 분석-개선 사이클을 통해 최종 95%를 초과하는 정확도를 달성하였습니다."
    )

    add_styled_paragraph(
        doc,
        "정확도 추이: 35% → 64% → 70% → 72% → (일시 하락) → 95%+",
        bold=True, color=COLOR_PRIMARY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(12), space_after=Pt(12)
    )

    add_styled_paragraph(doc, "주요 개선 단계:", bold=True)

    iter_items = [
        ("초기 (35%): ", "기본 ASR 정렬만으로는 한국어 텍스트의 특성상 오정렬이 다수 발생"),
        ("1차 개선 (64%): ", "매칭 임계값 조정 및 ASR 프롬프팅 전략 도입으로 대폭 향상"),
        ("2차 개선 (70%): ", "대형 ASR 모델을 활용한 2차 검증 도입으로 추가 복구"),
        ("3차 개선 (72%): ", "매칭 확인 알고리즘 도입, 이 과정에서 단일 변수 변경 원칙 확립"),
        ("최종 (95%+): ", "오디오 경계 처리 최적화 및 메타데이터 관리 개선으로 목표 달성"),
    ]
    for prefix, text in iter_items:
        add_bullet(doc, text, bold_prefix=prefix)

    add_styled_paragraph(
        doc,
        "각 iteration에서 실패 유형을 체계적으로 분류(정렬 오류, 경계 잡음, ASR 인식 한계 등)하고, "
        "지배적 실패 원인에 집중하여 개선하는 전략이 효과적이었습니다.",
        space_before=Pt(8)
    )

    # 3.6 Final Metrics
    add_heading(doc, "3.6 최종 품질 달성 현황 (Final Quality Results)", level=2)

    req_table = doc.add_table(rows=7, cols=3)
    req_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    req_table.style = "Table Grid"
    req_headers = ["요구사항", "달성 수준", "판정"]
    for i, h in enumerate(req_headers):
        cell = req_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=FONT_SIZE_TABLE, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, COLOR_TABLE_HEADER)

    req_data = [
        ["R1: 정렬 정확도", "95% 초과", "PASS"],
        ["R2: 경계 품질", "100%", "PASS"],
        ["R3: 종합 통과율", "95% 초과", "PASS"],
        ["R4: 메타데이터 무결성", "99% 이상", "PASS"],
        ["R5: 재현성", "확인됨", "PASS"],
        ["R6: 오디오 규격", "99% 이상", "PASS"],
    ]
    for r_idx, row_data in enumerate(req_data):
        for c_idx, cell_text in enumerate(row_data):
            cell = req_table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            if c_idx == 2:
                set_run_font(run, size=FONT_SIZE_TABLE, bold=True, color=COLOR_PASS)
            else:
                set_run_font(run, size=FONT_SIZE_TABLE)
            if r_idx % 2 == 1:
                set_cell_shading(cell, COLOR_TABLE_ALT)

    doc.add_paragraph()

    add_styled_paragraph(
        doc,
        "6개 요구사항 전체 충족 — 데이터셋 품질 기준 달성 완료",
        bold=True, color=COLOR_ACCENT,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(4), space_after=Pt(12)
    )


def section_4_final_dataset(doc):
    add_heading(doc, "4. 최종 데이터셋 (Final Dataset)", level=1)

    add_heading(doc, "4.1 데이터셋 규모 (Dataset Scale)", level=2)

    create_table(doc,
                 ["항목", "수치"],
                 [
                     ["총 문장 수", "약 4,200문장"],
                     ["총 오디오 시간", "약 8시간"],
                     ["오디오 포맷", "고해상도 무손실 WAV (Mono)"],
                     ["정렬 정확도", "95% 초과"],
                     ["품질 검증 완료", "6개 항목 전체 PASS"],
                 ])

    add_heading(doc, "4.2 데이터 품질 특성 (Quality Characteristics)", level=2)

    quality_items = [
        ("높은 정렬 정확도: ",
         "ASR 재검증 기준 95%를 초과하는 텍스트-오디오 정렬 정확도 달성"),
        ("깨끗한 오디오 경계: ",
         "인접 문장 bleeding 0%, 클릭/팝 아티팩트 제거 완료"),
        ("표준화된 오디오 구조: ",
         "모든 파일이 일관된 무음-음성-무음 구조를 갖추어 TTS 학습에 최적화"),
        ("다양한 문체 포괄: ",
         "문학, 기술, 교양, 내러티브 등 5개 도메인의 텍스트 포함"),
        ("전문 성우 녹음: ",
         "일관된 음색, 발화 속도, 감정 표현으로 높은 녹음 품질"),
    ]
    for prefix, text in quality_items:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading(doc, "4.3 잔여 실패 분석 (Residual Failure Analysis)", level=2)
    add_styled_paragraph(
        doc,
        "약 5%의 미통과 세그먼트 중, 절대 다수(95% 이상)가 ASR 모델의 한국어 인식 한계에 "
        "기인합니다. 이는 실제 오디오-텍스트 정렬이 올바르지만, 검증용 ASR 모델이 "
        "비일상적 어휘(문학적 표현, 전문 용어 등)를 정확히 재현하지 못하는 경우입니다."
    )
    add_styled_paragraph(
        doc,
        "즉, 잔여 실패의 대부분은 데이터 자체의 문제가 아닌 검증 수단의 한계로, "
        "실질적인 데이터셋 품질은 수치 이상으로 높습니다.",
        italic=True, color=COLOR_GRAY
    )


def section_5_model_training(doc):
    add_heading(doc, "5. 모델 학습 (Model Training)", level=1)

    add_heading(doc, "5.1 접근 방식 (Approach)", level=2)
    add_styled_paragraph(
        doc,
        "최신 flow-matching 기반 TTS 아키텍처를 한국어로 fine-tuning하는 전이학습(transfer learning) "
        "접근법을 채택하였습니다. 대규모 다국어 사전학습 모델의 음성 생성 능력을 활용하면서, "
        "한국어에 특화된 추가 학습을 수행합니다."
    )

    approach_items = [
        ("전이학습: ", "대규모 사전학습 TTS 모델을 기반으로 한국어 fine-tuning"),
        ("Vocabulary 확장: ", "사전학습 모델의 기존 토큰에 한국어 고유 문자를 추가하여 "
         "사전학습 지식을 보존하면서 한국어 지원을 확장"),
        ("임베딩 확장: ", "기존 임베딩 레이어를 보존한 채 한국어 토큰용 임베딩만 추가"),
        ("클라우드 GPU 활용: ", "고성능 클라우드 GPU를 on-demand로 활용하여 비용 효율적으로 학습"),
    ]
    for prefix, text in approach_items:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading(doc, "5.2 학습 전략 (Training Strategy)", level=2)
    add_styled_paragraph(
        doc,
        "2단계(Phase) 학습 전략을 채택하였습니다."
    )

    create_table(doc,
                 ["단계", "목표", "상태"],
                 [
                     ["Phase 1", "기본 한국어 음성 생성 능력 확보", "완료"],
                     ["Phase 2", "프로소디 및 자연스러움 향상을 위한 추가 학습", "계획 중"],
                 ])

    add_styled_paragraph(
        doc,
        "Phase 1에서는 한국어 발화의 기본적인 음소-음향 매핑을 학습하는 것에 집중하였으며, "
        "성공적으로 완료되었습니다. Phase 2에서는 추가 학습을 통해 "
        "프로소디(prosody), 자연스러움, 감정 표현력 등의 품질 향상을 목표로 합니다.",
        space_before=Pt(4)
    )

    add_heading(doc, "5.3 학습 효율 최적화 (Training Optimization)", level=2)
    add_styled_paragraph(
        doc,
        "제한된 GPU 자원에서 최대 효율을 달성하기 위해 다양한 최적화를 수행하였습니다."
    )

    opt_items = [
        ("Mixed Precision 학습: ", "메모리 사용량 절감 및 학습 속도 향상"),
        ("동적 배칭(Dynamic Batching): ", "오디오 길이 기반 프레임 단위 배칭으로 GPU 활용률 극대화"),
        ("Batch size 최적화: ", "GPU 메모리 활용률을 80% 이상으로 끌어올려 처리량 대폭 향상"),
        ("Gradient Accumulation: ", "실질적 배치 크기를 늘려 학습 안정성 확보"),
        ("체크포인트 자동 백업: ", "클라우드 환경의 데이터 유실 방지를 위한 실시간 백업 시스템"),
    ]
    for prefix, text in opt_items:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading(doc, "5.4 비용 효율성 (Cost Efficiency)", level=2)
    add_styled_paragraph(
        doc,
        "클라우드 GPU의 on-demand 특성을 활용하여, Phase 1 전체 학습을 "
        "매우 낮은 비용으로 완료하였습니다. 배치 크기 최적화를 통해 "
        "동일 학습량 대비 약 33%의 비용을 절감하였습니다."
    )

    create_table(doc,
                 ["항목", "결과"],
                 [
                     ["Phase 1 학습 시간", "수 시간 이내"],
                     ["배치 최적화 효과", "처리량 약 38% 향상, 학습 시간 약 33% 단축"],
                     ["비용 수준", "소규모 연구 예산으로 충분히 감당 가능"],
                 ])


def section_6_inference(doc):
    add_heading(doc, "6. 추론 및 활용 (Inference & Application)", level=1)

    add_heading(doc, "6.1 지원 인터페이스 (Supported Interfaces)", level=2)

    create_table(doc,
                 ["인터페이스", "설명", "적합 용도"],
                 [
                     ["CLI", "명령줄 기반 음성 생성", "배치 처리, 자동화"],
                     ["Web UI", "브라우저 기반 대화형 인터페이스", "데모, 테스트"],
                     ["Python API", "프로그래밍 인터페이스", "애플리케이션 통합"],
                 ])

    add_heading(doc, "6.2 주요 기능 (Key Features)", level=2)

    features = [
        ("텍스트 → 음성 변환: ", "한국어 텍스트를 입력하면 자연스러운 음성을 생성"),
        ("참조 음성 활용: ", "특정 참조 음성의 스타일을 반영한 음성 생성 가능"),
        ("배치 생성: ", "다수의 문장을 파일로 입력하여 일괄 생성"),
        ("모델 경량화: ", "학습 체크포인트에서 추론에 필요한 가중치만 추출하여 "
         "용량을 약 75% 절감"),
    ]
    for prefix, text in features:
        add_bullet(doc, text, bold_prefix=prefix)


def section_7_challenges(doc):
    add_heading(doc, "7. 한국어 TTS 구축의 주요 과제 (Key Challenges)", level=1)

    add_styled_paragraph(
        doc,
        "본 연구를 통해 한국어 TTS 데이터셋 구축 및 모델 학습 과정에서 "
        "발견된 주요 기술적 과제들을 정리하였습니다."
    )

    add_heading(doc, "7.1 한국어 음성 인식의 고유한 어려움", level=2)

    kr_items = [
        ("음운 변동의 복잡성: ",
         "한국어는 연음, 경음화, 비음화, 구개음화 등 다양한 음운 변동 규칙이 있어, "
         "실제 발화 음향과 표기 텍스트 사이의 괴리가 큽니다. "
         "이로 인해 ASR 모델의 전사 정확도가 영어 대비 낮아질 수 있습니다."),
        ("조사/어미의 유사성: ",
         "한국어 문장들은 공통된 조사와 어미를 공유하여, "
         "서로 다른 문장 간 텍스트 유사도의 기준선(baseline)이 높습니다. "
         "이는 자동 정렬 시 오탐(false positive)의 원인이 됩니다."),
        ("ASR 디코딩 전략의 중요성: ",
         "한국어와 같이 음향적 모호성이 높은 언어에서는 ASR 모델의 디코딩 전략 설정이 "
         "결과에 극적인 영향을 미칩니다. 특정 설정은 인식률을 치명적으로 저하시킬 수 있어 "
         "언어별 최적화가 필수적입니다."),
        ("전문 어휘의 인식 한계: ",
         "문학적 표현, IT 전문 용어, 외래어 혼용 등 비일상적 어휘에 대한 "
         "ASR 모델의 인식률이 현저히 낮아, 잔여 오류의 주된 원인이 됩니다."),
    ]
    for prefix, text in kr_items:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading(doc, "7.2 데이터 정제 과정의 교훈", level=2)

    cleanse_items = [
        ("단일 변수 변경 원칙: ",
         "여러 파라미터를 동시에 변경하면 개선/악화 원인을 특정할 수 없습니다. "
         "반드시 한 번에 하나의 변수만 조정하는 원칙이 중요합니다."),
        ("자동 메트릭의 한계: ",
         "자동화된 품질 측정만으로는 인지적(perceptual) 품질을 완전히 보장할 수 없으며, "
         "주기적인 수동 청취 검증이 필요합니다."),
        ("파라미터 이력 관리: ",
         "반복적 R&D 과정에서 이전에 revert된 설정이 다시 적용되는 "
         "regression이 발생할 수 있어 체계적인 변경 이력 관리가 필수적입니다."),
        ("정렬 흐름의 보존: ",
         "정렬 과정 중간에 추가 검증을 삽입하면 오히려 정렬 흐름을 방해하여 "
         "수량과 품질 모두 저하될 수 있습니다. 정렬과 검증은 분리하는 것이 효과적입니다."),
    ]
    for prefix, text in cleanse_items:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading(doc, "7.3 모델 학습의 교훈", level=2)

    train_items = [
        ("Vocabulary 설계: ",
         "다국어 사전학습 모델에 새 언어를 추가할 때, 기존 vocab을 유지하면서 "
         "신규 토큰을 병합하는 방식이 사전학습 지식 보존에 핵심적입니다."),
        ("자원 최적화의 중요성: ",
         "배치 크기, mixed precision 등의 최적화만으로도 "
         "학습 시간과 비용을 30% 이상 절감할 수 있습니다."),
        ("체크포인트 안전 관리: ",
         "클라우드 GPU 환경의 휘발성을 고려하여, 자동 백업 시스템 구축이 필수적입니다."),
    ]
    for prefix, text in train_items:
        add_bullet(doc, text, bold_prefix=prefix)


def section_8_next_steps(doc):
    add_heading(doc, "8. 향후 계획 (Next Steps & Roadmap)", level=1)

    roadmap = [
        ["단기", "Phase 2 추가 학습 완료", "프로소디 및 자연스러움 향상"],
        ["단기", "모델 품질 평가", "MOS 청취 테스트, 문체별 성능 비교"],
        ["중기", "데이터셋 확장", "미녹음 구간 추가 녹음, 데이터 증강"],
        ["중기", "서비스 통합", "API 서버 구축, 실시간 합성 최적화"],
        ["장기", "다화자 확장", "추가 화자 녹음으로 multi-speaker TTS 탐색"],
        ["장기", "프로덕션 배포", "모니터링 시스템 구축, 경량화 및 최적화"],
    ]
    create_table(doc,
                 ["시기", "항목", "설명"],
                 roadmap)

    add_styled_paragraph(
        doc,
        "본 연구를 통해 구축된 데이터 정제 파이프라인과 학습 인프라는 "
        "향후 다른 언어나 화자로의 확장에도 재사용 가능한 범용적 프레임워크입니다.",
        space_before=Pt(12), bold=True
    )


def section_9_conclusion(doc):
    add_heading(doc, "9. 결론 (Conclusion)", level=1)

    add_styled_paragraph(
        doc,
        "본 프로젝트는 한국어 TTS 음성합성 시스템의 End-to-End 구축을 목표로, "
        "데이터 수집부터 정제, 모델 학습, 추론까지의 전 과정을 수행하였습니다."
    )

    add_styled_paragraph(doc, "주요 성과:", bold=True, space_before=Pt(12))

    results = [
        ("자동화된 데이터 정제: ",
         "ASR 기반 6-Stage 반복 파이프라인을 통해 약 4,200문장, 8시간 분량의 "
         "고품질 한국어 TTS 데이터셋을 자동으로 구축하였습니다."),
        ("높은 품질 달성: ",
         "6개 품질 요구사항 전체를 충족하며, 정렬 정확도 95% 이상을 달성하였습니다. "
         "초기 35%에서 반복적 개선을 통해 도달한 수치입니다."),
        ("비용 효율적 학습: ",
         "클라우드 GPU를 활용한 비용 효율적 학습 체계를 구축하여, "
         "소규모 예산으로도 최신 TTS 모델의 한국어 fine-tuning이 가능함을 입증하였습니다."),
        ("한국어 특화 인사이트: ",
         "한국어의 음운 변동, ASR 디코딩 전략, 텍스트 유사도 특성 등 "
         "한국어 TTS 구축에 특화된 기술적 인사이트를 도출하였습니다."),
        ("재사용 가능한 프레임워크: ",
         "데이터 정제 파이프라인과 학습 인프라는 다른 언어/화자로의 확장에도 "
         "활용 가능한 범용 프레임워크로 설계되었습니다."),
    ]
    for prefix, text in results:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_paragraph()
    add_styled_paragraph(
        doc,
        "━" * 30,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        color=COLOR_GRAY
    )
    add_styled_paragraph(
        doc,
        "본 문서는 공개용 요약본입니다. 세부 기술 사양 및 파라미터는 비공개입니다.",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True, color=COLOR_GRAY, size=FONT_SIZE_SMALL
    )


# ══════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════

def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_KOREAN
    font.size = FONT_SIZE_BODY
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KOREAN)

    # Heading styles
    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = FONT_KOREAN
        heading_style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KOREAN)
        heading_style.font.color.rgb = COLOR_HEADING
        if level == 1:
            heading_style.font.size = Pt(18)
        elif level == 2:
            heading_style.font.size = Pt(14)
        else:
            heading_style.font.size = Pt(12)

    # Build document
    create_cover_page(doc)
    section_1_overview(doc)
    doc.add_page_break()
    section_2_data_collection(doc)
    doc.add_page_break()
    section_3_cleansing_pipeline(doc)
    doc.add_page_break()
    section_4_final_dataset(doc)
    doc.add_page_break()
    section_5_model_training(doc)
    doc.add_page_break()
    section_6_inference(doc)
    doc.add_page_break()
    section_7_challenges(doc)
    doc.add_page_break()
    section_8_next_steps(doc)
    doc.add_page_break()
    section_9_conclusion(doc)

    # Footer with page numbers
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)

    # Save
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Public report generated: {OUTPUT}")
    print(f"File size: {OUTPUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
