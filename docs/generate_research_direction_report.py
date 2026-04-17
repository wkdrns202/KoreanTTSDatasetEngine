# -*- coding: utf-8 -*-
"""
연구 방향 보고서 생성: SAM Data Engine 패러다임 기반 End-to-End TTS 구축 시스템 로드맵
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Style ──
style = doc.styles['Normal']
font = style.font
font.name = 'Malgun Gothic'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.3

rpr = style.element.get_or_add_rPr()
rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Malgun Gothic"/>')
rpr.append(rFonts)

for lvl in range(1, 4):
    hs = doc.styles[f'Heading {lvl}']
    hs.font.name = 'Malgun Gothic'
    hs.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)
    hrpr = hs.element.get_or_add_rPr()
    hrf = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Malgun Gothic"/>')
    hrpr.append(hrf)

def shade(cell, color):
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(s)

def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        shade(c, "1F3A6E")
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255,255,255)
                r.font.size = Pt(9.5)
                r.font.name = 'Malgun Gothic'
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    r.font.name = 'Malgun Gothic'
            if ri % 2 == 1:
                shade(c, "DAE3F3")
    if widths:
        for row_obj in t.rows:
            for ci, w in enumerate(widths):
                row_obj.cells[ci].width = Cm(w)
    doc.add_paragraph()  # spacing
    return t

def bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.2 + level*0.8)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True; r.font.size = Pt(10); r.font.name = 'Malgun Gothic'
        r2 = p.add_run(text)
        r2.font.size = Pt(10); r2.font.name = 'Malgun Gothic'
    else:
        r = p.add_run(text)
        r.font.size = Pt(10); r.font.name = 'Malgun Gothic'

def para(text):
    p = doc.add_paragraph(text)
    return p

def bold_para(bold_text, normal_text):
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True; r.font.name = 'Malgun Gothic'
    r2 = p.add_run(normal_text)
    r2.font.name = 'Malgun Gothic'
    return p

# ════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph()

t1 = doc.add_paragraph()
t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t1.add_run("SAM Data Engine 패러다임 기반\nEnd-to-End 한국어 TTS 구축 시스템\n연구 방향 보고서")
r.font.size = Pt(22); r.bold = True; r.font.name = 'Malgun Gothic'
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

doc.add_paragraph()
t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t2.add_run(
    "Segment Anything(Meta, 2023)의 Data Annotation 자동화 패러다임을\n"
    "STT/TTS 음성 파이프라인 구축 및 모델 학습으로 확장하는 연구 로드맵"
)
r.font.size = Pt(12); r.font.name = 'Malgun Gothic'
r.font.color.rgb = RGBColor(100,100,100)

doc.add_paragraph()
doc.add_paragraph()
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = m.add_run("2026년 3월 24일\n천영재 (테크 7기 박사과정)\n\n지도교수 코멘트 반영")
r.font.size = Pt(11); r.font.name = 'Malgun Gothic'

doc.add_page_break()

# ════════════════════════════════════════════════
# TOC
# ════════════════════════════════════════════════
doc.add_heading('목차', level=1)
toc = [
    "1. 배경: 지도교수 코멘트 및 연구 맥락",
    "2. 현재까지의 성과 요약",
    "3. SAM과 본 프로젝트의 구조적 동형성",
    "4. 연구 확장 방향: End-to-End TTS 구축 시스템",
    "   4.1 전체 시스템 아키텍처",
    "   4.2 Phase A — Data Engine 고도화",
    "   4.3 Phase B — 모델 학습 파이프라인 확립",
    "   4.4 Phase C — 추론/배포 시스템 구축",
    "   4.5 Phase D — Data Engine ↔ Model 피드백 루프",
    "5. 단계별 실행 로드맵",
    "6. 예상 기여 및 학술적 의의",
    "7. 리스크 및 완화 전략",
    "8. 결론",
    "참고문헌",
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════
# 1. 배경
# ════════════════════════════════════════════════
doc.add_heading('1. 배경: 지도교수 코멘트 및 연구 맥락', level=1)

doc.add_heading('1.1 지도교수 코멘트', level=2)

# Quote box
qt = doc.add_paragraph()
qt.paragraph_format.left_indent = Cm(1.0)
qt.paragraph_format.right_indent = Cm(1.0)
qt.paragraph_format.space_before = Pt(8)
qt.paragraph_format.space_after = Pt(8)
r = qt.add_run(
    '"Meta의 Segment Anything(SAM)의 주요 contribution 중 하나가 data annotation 자동화 '
    '파이프라인인데, 논문을 참고하여 STT/TTS 용 파이프라인 구축 + 모델 학습으로의 확장으로 '
    '연구를 이어가면 좋을 것 같네요."'
)
r.italic = True; r.font.size = Pt(11); r.font.name = 'Malgun Gothic'
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

doc.add_heading('1.2 코멘트 해석 및 연구 방향 도출', level=2)

para(
    '지도교수의 코멘트는 세 가지 핵심 방향을 제시한다.'
)
bullet(' SAM의 Data Engine(어노테이션 자동화 → 데이터 품질 향상 → 모델 개선의 선순환) '
       '패러다임을 음성 도메인에 체계적으로 적용할 것', '(1) SAM 패러다임의 도메인 전이:')
bullet(' 현재 데이터 정제(Cleanse) 단계에 국한된 연구를 STT(음성인식)/TTS(음성합성) '
       '모델 학습까지 확장할 것', '(2) 파이프라인 → 모델 학습으로의 확장:')
bullet(' 데이터 수집 → 정제 → 학습 → 추론/배포까지 전 과정을 통합하는 '
       'End-to-End 시스템을 구축할 것', '(3) End-to-End 시스템 관점:')

doc.add_heading('1.3 연구의 궁극적 목표', level=2)
bold_para('Goal: ', '데이터 수집(Record)부터 추론/배포(Inference/Deploy)까지를 포괄하는 '
    'End-to-End 한국어 TTS 구축 시스템을 확립하고, SAM의 Data Engine 패러다임을 '
    '음성 도메인에 적용하여 데이터-모델 간 선순환 구조를 구현한다.')

# ════════════════════════════════════════════════
# 2. 현재까지의 성과
# ════════════════════════════════════════════════
doc.add_heading('2. 현재까지의 성과 요약', level=1)

make_table(
    ["영역", "달성 사항", "상태"],
    [
        ["데이터 수집", "스튜디오 녹음, 5개 대본, 수천 문장 규모 확보", "완료"],
        ["데이터 정제\n(Data Engine v1)", "ASR 기반 6-Stage 반복 파이프라인 구축\n정렬 정확도 35% → 95%+ (6회 iteration)", "완료"],
        ["품질 보증", "6개 요구사항(R1~R6) 전체 PASS\n4,200문장/8시간 데이터셋 확정", "완료"],
        ["모델 학습\n(Phase 1)", "F5TTS 기본 한국어 음성 생성 학습\nMixed Precision, Dynamic Batching 최적화", "1회 완료"],
        ["모델 학습\n(Phase 2)", "프로소디/감정 표현 향상을 위한 추가 학습", "대본 작성 완료\n학습 준비 중"],
        ["추론/배포", "CLI, Web UI, API 인터페이스 설계", "계획 수립"],
    ],
    widths=[3.0, 7.5, 3.5]
)

para(
    '현재 전체 파이프라인 4단계(수집→정제→학습→추론) 중, 수집과 정제가 완료되었고 '
    '학습이 탐색적 수준에서 1회 수행된 상태이다. 지도교수 코멘트는 이 지점에서 '
    'SAM 패러다임을 참고하여 학습↔정제 간 피드백 루프를 확립하고, '
    '추론/배포까지 완결하라는 방향을 제시한다.'
)

# ════════════════════════════════════════════════
# 3. SAM과의 구조적 동형성
# ════════════════════════════════════════════════
doc.add_heading('3. SAM과 본 프로젝트의 구조적 동형성', level=1)

para(
    'SAM(Kirillov et al., 2023)은 이미지 세그멘테이션 분야의 Foundation Model로서, '
    '(1) Promptable Segmentation이라는 범용 태스크 정의, (2) Data Engine을 통한 '
    '어노테이션 자동화 및 데이터-모델 선순환, (3) Zero-shot Transfer 능력을 핵심 기여로 제시하였다. '
    '본 프로젝트는 도메인(이미지→음성)은 다르지만, 이와 구조적으로 동형(isomorphic)인 '
    '접근을 독립적으로 구현하였다.'
)

doc.add_heading('3.1 패러다임 대응 구조', level=2)

make_table(
    ["SAM (이미지 도메인)", "본 프로젝트 (음성 도메인)", "대응 관계"],
    [
        ["Promptable Segmentation\n프롬프트 → 유효 마스크", "Promptable Alignment\n대본 텍스트 → 음성 구간", "태스크 정의"],
        ["ViT-H Image Encoder\n(1회 실행, amortized)", "Whisper ASR Encoder\n(1회 전사, amortized)", "Foundation Model"],
        ["Data Engine 3단계\nManual→Semi-auto→Full-auto", "6-Stage 반복 파이프라인\n6회 iteration (35%→95%+)", "반복적 품질 향상"],
        ["IoU confidence + NMS\n+ Stability check", "유사도 임계값 + Tier 2 검증\n+ 6개 요구사항 체계", "품질 필터링"],
        ["SA-1B (1.1B 마스크)", "4,200문장 음성-텍스트 쌍\n데이터셋", "산출물"],
        ["Zero-shot Transfer\n(23개 데이터셋)", "재현성 검증\n(다른 데이터셋 적용, 계획)", "범용성"],
    ],
    widths=[4.5, 4.5, 3.5]
)

doc.add_heading('3.2 핵심 차이점과 독자성', level=2)

make_table(
    ["차원", "SAM", "본 프로젝트"],
    [
        ["모델 업데이트", "Foundation Model 재학습\n(6회 retrain)", "모델 고정(Whisper)\n알고리즘/파라미터 개선"],
        ["리소스", "256+ GPU 클러스터", "단일 RTX 3060 Ti (8GB)"],
        ["모호성 대응", "Multi-output (3개 마스크)", "단일 출력 + 도메인 지식 보완"],
        ["언어 특수성", "해당 없음", "한국어 음운 변동, 조사/어미\n유사성 등 언어 특화 대응"],
    ],
    widths=[3.5, 5.0, 5.0]
)

bold_para(
    '독자적 기여: ',
    'SAM이 대규모 리소스 기반 모델 재학습으로 Data Engine을 구동한 반면, '
    '본 프로젝트는 Foundation Model(Whisper)을 고정한 채 알고리즘 최적화만으로 '
    '동일한 패러다임을 성공적으로 구현하였다. 이는 제한된 리소스 환경에서도 '
    'Data Engine 접근이 유효함을 보여주는 사례이다.'
)

# ════════════════════════════════════════════════
# 4. 연구 확장 방향
# ════════════════════════════════════════════════
doc.add_heading('4. 연구 확장 방향: End-to-End TTS 구축 시스템', level=1)

para(
    '지도교수 코멘트에 따라, 현재의 데이터 정제 파이프라인을 "Data Engine v1"으로 정의하고, '
    'SAM의 Data Engine 패러다임을 참고하여 모델 학습까지 포함하는 End-to-End 시스템으로 '
    '확장한다. 전체 시스템은 4개 Phase로 구성된다.'
)

# 4.1
doc.add_heading('4.1 전체 시스템 아키텍처', level=2)

para(
    'SAM이 Task-Model-Data Engine 세 요소의 상호작용으로 구성되듯, '
    '본 시스템도 세 축의 상호작용으로 구성한다.'
)

make_table(
    ["SAM의 3요소", "본 시스템의 대응", "설명"],
    [
        ["Task\n(Promptable Segmentation)", "Task\n(Promptable Speech Alignment\n& Synthesis)", "대본 → 정렬된 음성 데이터 생성\n+ 텍스트 → 음성 합성"],
        ["Model\n(SAM)", "Model\n(Whisper + F5TTS)", "ASR(정제용) + TTS(합성용)\n두 모델의 협업 체계"],
        ["Data Engine\n(3-stage annotation)", "Data Engine v2\n(4-phase pipeline)", "수집→정제→학습→추론의\n전 과정을 순환하는 엔진"],
    ],
    widths=[4.0, 4.5, 5.5]
)

# 4.2
doc.add_heading('4.2 Phase A — Data Engine 고도화 (정제 파이프라인 발전)', level=2)

bold_para('목표: ', '현재 Data Engine v1(6-Stage 파이프라인)을 SAM의 방법론을 참고하여 고도화한다.')

doc.add_heading('A-1. Stability-based Filtering 도입', level=3)
para(
    'SAM은 mask probability를 ±δ 변동시켜 안정적인 마스크만 선별하는 stability check를 적용한다. '
    '이를 음성 정렬에 적용한다.'
)
bullet('매칭 결과가 similarity threshold ± 0.05에서도 동일 → "안정 매칭" (자동 확정)')
bullet('매칭 결과가 변동에 민감 → "불안정 매칭" (Tier 2 대형 모델 검증 대상)')
bold_para('기대 효과: ', 'Tier 2(Whisper large) 검증 대상을 선별하여, 8GB VRAM 제약 하에서의 비용 효율성 향상. '
    '현재 전체 실패 건에 대해 Tier 2를 적용하는 방식 대비 연산량 절감 기대.')

doc.add_heading('A-2. Multi-hypothesis Matching', level=3)
para(
    'SAM의 ambiguity-aware multi-output(3개 마스크 동시 출력) 전략을 참고하여, '
    '현재 greedy 단일 매칭 대신 상위 3개 후보를 유지하고, '
    '전후 문맥의 monotonic alignment 제약으로 최종 선택하는 전략을 도입한다.'
)
bold_para('기대 효과: ', 're-sync 실패 감소, 연속 매칭 실패(CONSEC_FAIL_LIMIT) 발생 빈도 저하.')

doc.add_heading('A-3. 재현성 검증 (Zero-shot Transfer 대응)', level=3)
para(
    'SAM이 23개 데이터셋에서 zero-shot 성능을 입증했듯, 본 파이프라인도 '
    '다른 화자/언어/도메인의 데이터에 적용하여 범용성을 검증한다.'
)
bullet('동일 파이프라인으로 다른 한국어 화자의 녹음 데이터 처리')
bullet('영어/일본어 등 다른 언어의 녹음 데이터에 대한 적용 가능성 탐색')
bullet('파라미터를 언어별 플러그인으로 분리하는 모듈화 설계')

# 4.3
doc.add_heading('4.3 Phase B — 모델 학습 파이프라인 확립', level=2)

bold_para('목표: ', 'Data Engine이 생산한 데이터셋으로 TTS 모델을 체계적으로 학습하고, '
    '모델 품질을 정량적으로 평가하는 체계를 확립한다.')

doc.add_heading('B-1. Phase 1 학습 완성 (기본 음성 생성)', level=3)
bullet('현재 F5TTS Phase 1에서 발음 미출력 문제 → 추가 학습 데이터 투입으로 해결')
bullet('학습 파라미터 최적화: Dynamic Batching, Gradient Accumulation 설정 고도화')
bullet('체크포인트 관리: HuggingFace Hub 기반 버전 관리 체계 확립')

doc.add_heading('B-2. Phase 2 학습 (프로소디/감정 향상)', level=3)
bullet('감정/억양/캐릭터가 부여된 새 녹음 데이터 활용 (대본 작성 완료)')
bullet('데이터 라벨에 감정/스타일 태그 추가 → 조건부 생성(conditional generation) 학습')
bullet('프로소디 품질 측정을 위한 객관적 메트릭 도입 (F0 contour, duration variance 등)')

doc.add_heading('B-3. 모델 품질 평가 체계', level=3)
make_table(
    ["평가 방법", "측정 항목", "목적"],
    [
        ["MOS 청취 테스트", "자연스러움, 명료도, 감정 표현", "인지적 품질의 정량화"],
        ["자동 메트릭", "MCD, PESQ, F0 RMSE", "반복 가능한 객관적 비교"],
        ["A/B 테스트", "Phase 1 vs Phase 2,\n대본별 성능 비교", "학습 진전도 추적"],
        ["ASR 역검증", "합성 음성의 Whisper 전사 정확도", "발음 정확성 자동 평가"],
    ],
    widths=[3.5, 4.5, 5.5]
)

# 4.4
doc.add_heading('4.4 Phase C — 추론/배포 시스템 구축', level=2)

bold_para('목표: ', '학습된 TTS 모델을 실제 사용 가능한 서비스로 배포한다.')

make_table(
    ["인터페이스", "기능", "적합 용도", "기술 스택"],
    [
        ["CLI", "명령줄 기반 음성 생성", "배치 처리, 자동화", "Python argparse"],
        ["Web UI", "브라우저 기반 대화형 인터페이스", "데모, 테스트, 피드백 수집", "Gradio / Streamlit"],
        ["Python API", "프로그래밍 인터페이스", "애플리케이션 통합", "FastAPI"],
    ],
    widths=[2.5, 4.0, 3.5, 3.5]
)

bullet('모델 경량화: 학습 체크포인트에서 추론 가중치만 추출 (약 75% 용량 절감)')
bullet('실시간 합성 최적화: TorchScript/ONNX 변환으로 추론 속도 향상')
bullet('참조 음성 기반 스타일 전이: 특정 음성의 톤/감정을 반영한 합성')

# 4.5
doc.add_heading('4.5 Phase D — Data Engine ↔ Model 피드백 루프 (SAM 핵심 전략)', level=2)

bold_para('핵심: ', '이 Phase가 SAM 패러다임의 가장 핵심적인 적용이다. '
    'SAM에서 모델이 어노테이션을 보조하고, 새 데이터가 모델을 개선하는 선순환을 구현했듯, '
    'TTS 모델의 출력을 Data Engine에 피드백하여 전체 시스템을 발전시킨다.')

doc.add_heading('D-1. TTS 출력 기반 데이터 품질 역검증', level=3)
para(
    '학습된 TTS 모델로 텍스트를 합성한 후, 합성 음성을 다시 Whisper로 전사하여 '
    '원본 텍스트와 비교한다. 일치도가 낮은 항목은 학습 데이터의 정렬 오류일 가능성이 높으므로, '
    'Data Engine으로 되돌려 재검증한다.'
)
bullet('합성 음성 → ASR 전사 → 원본 대본과 비교 → 불일치 항목 플래깅')
bullet('플래깅된 항목을 Data Engine에서 재정렬 또는 quarantine 처리')
bold_para('SAM 대응: ', 'SAM의 Semi-automatic 단계에서 모델이 먼저 예측하고 사람이 보완하는 구조와 동일.')

doc.add_heading('D-2. 합성 데이터 증강 (Data Augmentation)', level=3)
para(
    'TTS 모델이 일정 품질에 도달하면, 합성 음성을 추가 학습 데이터로 활용하는 '
    'self-training 루프를 구축한다. '
    '이는 SAM의 Fully automatic 단계에서 모델이 스스로 1.1B 마스크를 생성한 것에 대응한다.'
)
bullet('고품질 합성 음성을 증강 데이터로 추가 → 모델 재학습')
bullet('다양한 프로소디/속도 변형 생성으로 데이터 다양성 확보')
bullet('합성 데이터 품질 게이팅: MOS 또는 자동 메트릭 기준 초과 항목만 사용')

doc.add_heading('D-3. 전체 선순환 구조', level=3)
para(
    '최종적으로 다음과 같은 선순환 구조가 완성된다.'
)

para(
    '녹음 데이터 수집 → [Data Engine] 자동 정제 → 학습용 데이터셋 → '
    '[TTS 모델 학습] → 합성 음성 생성 → [역검증] ASR 전사 비교 → '
    '정제 품질 피드백 → [Data Engine 개선] → 더 좋은 데이터셋 → 더 좋은 모델 → ...'
)

bold_para(
    'SAM과의 대응: ',
    'SAM: annotate → train → data → model → (loop). '
    '본 시스템: record → cleanse → train → synthesize → verify → (loop). '
    '두 시스템 모두 "모델이 데이터 품질을 개선하고, 개선된 데이터가 모델을 발전시키는" '
    '동일한 선순환 원리에 기반한다.'
)

# ════════════════════════════════════════════════
# 5. 실행 로드맵
# ════════════════════════════════════════════════
doc.add_heading('5. 단계별 실행 로드맵', level=1)

make_table(
    ["시기", "Phase", "핵심 과제", "산출물", "성공 기준"],
    [
        ["단기\n(1~2개월)", "A\nData Engine\n고도화",
         "Stability filtering 구현\nMulti-hypothesis 매칭\n재현성 검증 (다른 데이터셋)",
         "Data Engine v2 코드\n재현성 검증 보고서",
         "Tier 2 호출 50%↓\nre-sync 실패 30%↓\n타 데이터셋 정확도 90%+"],

        ["단기\n(1~2개월)", "B-1\n학습 완성",
         "Phase 1 추가 학습\nPhase 2 학습 실행\n평가 체계 구축",
         "Phase 2 체크포인트\nMOS 평가 결과",
         "MOS 3.5+ (5점 만점)\n발음 정확도 90%+"],

        ["중기\n(3~4개월)", "B-2 + C\n학습 고도화\n+ 배포",
         "조건부 생성 학습\nAPI 서버 구축\n모델 경량화",
         "감정 TTS 모델\nAPI 엔드포인트\n경량 추론 모델",
         "감정 표현 구분 가능\nAPI 응답 < 2초\n모델 크기 75%↓"],

        ["중기\n(3~4개월)", "D\n피드백 루프",
         "TTS 역검증 시스템\n합성 데이터 증강\n선순환 구조 확립",
         "역검증 파이프라인\n증강 데이터셋\n통합 시스템",
         "역검증으로 오류 5%↓\n증강 후 MOS 향상"],

        ["장기\n(6개월+)", "확장",
         "다화자 TTS 탐색\n다국어 파이프라인\n프로덕션 배포",
         "Multi-speaker 모델\n범용 프레임워크\n모니터링 시스템",
         "2+ 화자 지원\n2+ 언어 적용\n안정적 서비스 운영"],
    ],
    widths=[2.2, 2.5, 4.0, 3.0, 3.3]
)

# ════════════════════════════════════════════════
# 6. 기여 및 학술적 의의
# ════════════════════════════════════════════════
doc.add_heading('6. 예상 기여 및 학술적 의의', level=1)

doc.add_heading('6.1 학술적 기여', level=2)
bullet(' SAM의 Data Engine 패러다임이 이미지뿐 아니라 음성 도메인에서도 유효함을 실증. '
       'Foundation Model(Whisper)을 고정한 채 알고리즘 최적화만으로 Data Engine을 구동한 '
       '최초의 체계적 사례 보고.', '(1) Cross-domain Data Engine 패러다임 검증:')
bullet(' TTS 학습 결과를 데이터 정제에 피드백하는 양방향 선순환 구조의 제안 및 구현. '
       'SAM의 단방향(모델→데이터) 루프를 양방향(모델↔데이터)으로 확장.', '(2) TTS-Data Engine 피드백 루프:')
bullet(' 대규모 GPU 클러스터 없이도 단일 소비자용 GPU(RTX 3060 Ti, 8GB)로 '
       'End-to-End TTS 구축이 가능함을 실증.', '(3) 제한된 리소스 환경에서의 End-to-End TTS:')

doc.add_heading('6.2 실용적 기여', level=2)
bullet('한국어 TTS 구축의 비용/시간 장벽을 낮추는 오픈소스 파이프라인 제공')
bullet('녹음 → 배포까지의 전 과정을 자동화하여 음성합성 모델 R&D 사이클 단축')
bullet('다른 언어/화자로 확장 가능한 범용 프레임워크의 기반 확립')

doc.add_heading('6.3 논문 작성 시 SAM 인용 전략', level=2)
para(
    '본 연구는 SAM(Kirillov et al., 2023)이 이미지 세그멘테이션에서 제안한 '
    'Data Engine 패러다임 — "모델이 어노테이션을 보조하고, 새 데이터가 모델을 개선하는 '
    '선순환 구조" — 을 음성 도메인에 적용하여, STT/TTS 학습용 데이터 어노테이션 자동화 '
    '파이프라인을 구축하였다. SAM이 Manual→Semi-automatic→Fully automatic의 3단계 전이를 '
    '통해 1.1B 규모의 세그멘테이션 데이터셋을 구축했듯, 본 연구는 6회의 반복적 개선을 통해 '
    '정렬 정확도를 35%에서 95%+로 향상시키고, 나아가 TTS 모델 학습과의 피드백 루프를 '
    '통해 Data Engine의 범위를 데이터 정제 → 모델 학습 → 합성 검증까지 확장하였다.'
)

# ════════════════════════════════════════════════
# 7. 리스크
# ════════════════════════════════════════════════
doc.add_heading('7. 리스크 및 완화 전략', level=1)

make_table(
    ["리스크", "영향도", "완화 전략"],
    [
        ["GPU VRAM 부족\n(8GB 제약)", "높음",
         "Tier 2 선별적 적용 (Stability filtering)\n클라우드 GPU on-demand 활용\n모델 경량화 (quantization)"],
        ["TTS 모델 품질 부족\n(Phase 1 발음 미출력)", "높음",
         "추가 학습 데이터 투입\n학습 하이퍼파라미터 체계적 탐색\n다른 TTS 아키텍처 탐색 (병행)"],
        ["재현성 실패\n(다른 데이터셋 적용 시)", "중간",
         "언어별 파라미터 플러그인 분리\n적용 실패 원인 체계적 분석\n점진적 확장 (한국어 다른 화자 → 다른 언어)"],
        ["피드백 루프 발산\n(합성 데이터 품질 저하 전파)", "중간",
         "합성 데이터 품질 게이팅 (MOS 기준)\n합성/실제 데이터 비율 제한\n주기적 인간 검수 병행"],
        ["연구 기간 초과", "낮음",
         "Phase별 독립적 산출물 확보\n핵심 Phase(A, B-1)에 우선 집중\n장기 과제는 후속 연구로 분리"],
    ],
    widths=[3.5, 2.0, 8.0]
)

# ════════════════════════════════════════════════
# 8. 결론
# ════════════════════════════════════════════════
doc.add_heading('8. 결론', level=1)

para(
    '지도교수 코멘트에 따라, SAM의 Data Annotation 자동화 패러다임을 '
    'STT/TTS 음성 파이프라인 구축 및 모델 학습으로 확장하는 연구 방향을 수립하였다.'
)

para('본 연구의 확장 방향은 다음 세 가지로 요약된다.')

bullet(
    ' 현재의 6-Stage 반복 파이프라인(Data Engine v1)에 SAM의 stability filtering, '
    'multi-hypothesis matching 등을 적용하여 Data Engine v2로 고도화한다.',
    '(1) Data Engine 고도화:'
)
bullet(
    ' Phase 1/2 학습을 완성하고, 체계적인 평가 체계(MOS, 자동 메트릭)를 구축하여 '
    '모델 품질을 정량적으로 추적한다.',
    '(2) 모델 학습 파이프라인 확립:'
)
bullet(
    ' TTS 모델의 합성 결과를 ASR로 역검증하여 데이터 품질을 개선하고, '
    '합성 데이터 증강으로 모델을 다시 발전시키는 SAM식 선순환 구조를 구현한다.',
    '(3) Data Engine ↔ Model 피드백 루프:'
)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run(
    '궁극적으로, 데이터 수집(Record)부터 추론/배포(Inference/Deploy)까지를 포괄하는 '
    'End-to-End 한국어 TTS 구축 시스템을 확립하고, SAM이 이미지 도메인에서 입증한 '
    '"Data Engine 패러다임"이 음성 도메인에서도 유효함을 실증하는 것이 본 연구의 목표이다.'
)
r.font.size = Pt(11); r.font.name = 'Malgun Gothic'; r.bold = True

# ════════════════════════════════════════════════
# References
# ════════════════════════════════════════════════
doc.add_heading('참고문헌', level=1)
refs = [
    'Kirillov, A., Mintun, E., Ravi, N., Mao, H., Rolland, C., Gustafson, L., Xiao, T., '
    'Whitehead, S., Berg, A.C., Lo, W.-Y., Dollar, P., & Girshick, R. (2023). '
    'Segment Anything. Meta AI Research, FAIR.',

    '천영재 (2026). 인공지능 음성합성(TTS) 파이프라인 구축: 데이터 수집 및 라벨링 자동화 '
    '프로세스 연구개발 요약본. 연구 진행 보고서.',

    'Radford, A., Kim, J.W., Xu, T., Brockman, G., McLeavey, C., & Sutskever, I. (2023). '
    'Robust Speech Recognition via Large-Scale Weak Supervision. Proc. ICML 2023.',

    'Bommasani, R. et al. (2021). On the Opportunities and Risks of Foundation Models. '
    'arXiv:2108.07258.',

    'Shirahata, Y., Park, B., Yamamoto, R., & Tachibana, K. (2024). Audio-conditioned '
    'phonemic and prosodic annotation for building text-to-speech models from unlabeled speech '
    'data. Proc. Interspeech 2024, 2795-2799.',

    'Zalkow, F., Govalkar, P., Mueller, M., Habets, E.A.P., & Dittmar, C. (2023). '
    'Evaluating Speech-Phoneme Alignment and Its Impact on Neural Text-To-Speech Synthesis. '
    'Proc. IEEE ICASSP 2023.',

    'Lee, K.-N. & Chung, M. (2007). Morpheme-Based Modeling of Pronunciation Variation '
    'for Large Vocabulary Continuous Speech Recognition in Korean. IEICE Transactions on '
    'Information and Systems, E90-D(7), 1063-1072.',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(f'[{i}] {ref}')
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(9.5)

# ── Save ──
out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                    'SAM_TTS_Research_Direction_Report.docx')
doc.save(out)
print(f'Saved: {out}')
